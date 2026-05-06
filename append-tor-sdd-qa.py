#!/usr/bin/env python3
"""Append ToR and SDD Q&As to existing QA document"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document('/home/z/my-project/ba-practice/Interview_Practice_BA_Portfolio_QA.docx')

DARK = RGBColor(26, 26, 26)
GRAY_ANSWER = RGBColor(44, 44, 44)

def add_qa_block(doc, q_number, question, answer_paragraphs):
    """Add a Q&A block matching the existing document style"""
    
    # Question line with bottom border
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(14)
    
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '1',
        qn('w:space'): '1', qn('w:color'): 'DDDDDD'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    run_q = p.add_run(f'Q{q_number}')
    run_q.bold = True
    run_q.font.size = Pt(11)
    run_q.font.color.rgb = DARK
    run_q.font.name = 'Calibri'
    
    run_text = p.add_run(f'  {question}')
    run_text.font.size = Pt(11)
    run_text.font.color.rgb = DARK
    run_text.font.name = 'Calibri'
    
    # Answer paragraphs
    for para_text in answer_paragraphs:
        a = doc.add_paragraph()
        a.paragraph_format.space_after = Pt(6)
        a.paragraph_format.line_spacing = Pt(14)
        run = a.add_run(para_text)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY_ANSWER
        run.font.name = 'Calibri'

# ======================================================================
# NEW Q&A ENTRIES - ToR and SDD
# ======================================================================

# Section divider
divider = doc.add_paragraph()
divider.paragraph_format.space_before = Pt(24)
divider.paragraph_format.space_after = Pt(12)
pPr = divider._p.get_or_add_pPr()
pBdr = pPr.makeelement(qn('w:pBdr'), {})
bottom = pBdr.makeelement(qn('w:bottom'), {
    qn('w:val'): 'single', qn('w:sz'): '4',
    qn('w:space'): '1', qn('w:color'): '005A9C'
})
pBdr.append(bottom)
pPr.append(pBdr)

section_label = doc.add_paragraph()
section_label.paragraph_format.space_after = Pt(12)
run = section_label.add_run('Terms of Reference (ToR) and Service Design Document (SDD) - LMAS Project')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 90, 156)
run.font.name = 'Calibri'

# ─── Q42: Why do we need a ToR? ───
add_qa_block(doc, 42, 
    'In a BA project, you already have a BRD that describes the business requirements and scope. Why do you need a separate Terms of Reference document? What does the ToR cover that the BRD does not?',
    [
        'This is a very practical question, and the answer explains a common confusion. The BRD and the ToR serve completely different purposes, and you need both.',
        
        'The BRD answers the question: "What does the business need?" It describes the business goals, the stakeholder requirements, the functional and non-functional requirements, and the success criteria. The BRD is written from the business perspective. When you read a BRD, you understand what problem we are solving and what the expected outcome looks like.',
        
        'The ToR answers a different question: "How will this project be organized and governed?" The ToR defines the project boundaries, who is responsible for what, what the timeline is, what the governance structure looks like, and what the constraints are. The ToR is written from the project management perspective.',
        
        'Let me give you a concrete example. In the LMAS project for the State Employment Agency, the BRD says: "Citizens must be able to submit employment service applications through digital channels." This is a business requirement. The ToR says: "The BA Lead is responsible for writing the FRD by the end of Phase 2, and the Product Owner must approve it before development starts." This is a project governance rule.',
        
        'Here is a simple way to think about it. The BRD defines the "what" and the "why." The ToR defines the "who, when and how." Without a BRD, the team does not know what to build. Without a ToR, the team does not know who decides what, when things are due, and how decisions are made.',
        
        'In government projects like LMAS, the ToR is especially important because multiple government organizations are involved. The ToR makes it clear which agency provides data, who approves requirements, who signs off on deliverables and how disagreements are resolved. Without this clarity, a multi-agency project can easily get stuck.',
    ]
)

# ─── Q43: What should a ToR contain? ───
add_qa_block(doc, 43,
    'What is the correct structure of a Terms of Reference? What sections should it contain and in what order?',
    [
        'A ToR has a standard structure that most government and enterprise projects follow. In my LMAS ToR, I used ten main sections. Let me walk through each one.',
        
        'Section 1 is Introduction. This includes the document purpose, the background of why the project exists, and definitions of abbreviations. The introduction sets the context for everyone who reads the document. It answers: why are we doing this project?',
        
        'Section 2 is Project Objectives. This is where you define the strategic goal and the specific measurable objectives. Each objective has an ID, a name and a description. For example, OBJ-05 in my LMAS ToR says: "Reduce average service processing time from 14 calendar days to 5 business days." You also include success criteria with specific targets here.',
        
        'Section 3 is Project Scope. This is one of the most important sections. It has four parts: in-scope items, out-of-scope items, assumptions and constraints. The in-scope list defines what the project will deliver. The out-of-scope list is equally important because it prevents scope creep. Assumptions are things you believe to be true but have not verified yet. Constraints are limitations you cannot change, like budget limits or regulatory requirements.',
        
        'Section 4 is Deliverables. This is a complete list of every document, system and report that the project will produce. Each deliverable has an ID, a description, the phase when it is due and the person responsible. In my ToR, I listed 14 deliverables: 8 business analysis deliverables like ToR, BRD, SDD, FRD, SRS, API Specifications, BPMN models and UAT Plan, and 6 technical deliverables like the system architecture, database schema, Telegram bot and web portal.',
        
        'Section 5 is Project Timeline. This shows the phased approach with milestones. In my LMAS ToR, I defined 4 phases across 12 months with 8 key milestones. Each milestone has a date and a description of what "done" looks like.',
        
        'Section 6 is Team Structure and Roles. This defines every role in the project, which organization they come from and what their responsibilities are. It also includes the governance structure: steering committee meetings, sprint reviews and stakeholder syncs with defined frequency and participants.',
        
        'Section 7 is Stakeholder Communication. This is a detailed communication plan that says who you communicate with, through what channel, how often and about what topics. For example, "Steering Committee meets monthly to review project status, risks and budget."',
        
        'Section 8 is Risk Management. This includes a risk register with probability, impact and mitigation for each risk, plus an escalation path for when risks become critical issues.',
        
        'Section 9 is Quality Standards. This defines documentation standards, development standards and acceptance criteria for deliverables.',
        
        'Section 10 is Approval. This includes a signature table where each stakeholder signs to confirm their agreement with the ToR, plus a revision history table.',
    ]
)

# ─── Q44: Why do we need an SDD when we have BRD? ───
add_qa_block(doc, 44,
    'You already have a BRD that covers business requirements and scope. You also have a Gap Analysis that covers As-Is and To-Be processes. Why do you need a separate Service Design Document? What does the SDD cover that the BRD and Gap Analysis do not?',
    [
        'This is the most important question about the SDD, because it is the document that people understand the least. Let me explain the difference clearly.',
        
        'The BRD answers: "What does the business need?" It talks about business goals, stakeholders, scope and requirements. The Gap Analysis answers: "What is wrong with the current process and what should the future process look like?" It shows the As-Is state and the To-Be state at the process level.',
        
        'The SDD answers a completely different question: "How should the citizen experience this service from start to finish?" The SDD takes the citizen perspective and designs the complete service experience, not just the process or the requirements.',
        
        'Let me explain with the LMAS example. The BRD says: "The system shall allow citizens to submit employment service applications through digital channels." The Gap Analysis shows the As-Is process: citizen visits office, waits in queue, submits paper form. And the To-Be process: citizen opens Telegram bot, fills digital form, system validates data automatically.',
        
        'But none of these documents answer the following questions: Who is the citizen using this service? What is their life situation? What are their pain points? What channels do they prefer? What happens at each interaction point? How do multiple government agencies coordinate behind the scenes to deliver this service? What SLA targets should we set? How do we measure if the service is good?',
        
        'These are exactly the questions the SDD answers. The SDD has several unique sections that do not exist in the BRD or Gap Analysis.',
        
        'First, User Research and Personas. The SDD starts by defining who the users are. In my LMAS SDD, I created three personas: a young job seeker who is comfortable with technology, a mid-career worker who recently lost their job and needs guidance, and an employer who wants to post vacancies online. The BRD mentions stakeholders, but it does not go deep into user profiles and behavioral patterns.',
        
        'Second, Service Journey Analysis. The Gap Analysis shows the As-Is and To-Be processes, but the SDD goes deeper. It maps the complete citizen journey including emotional states, pain points at each stage and improvement opportunities. The journey analysis is more detailed because it includes the citizen experience, not just the process steps.',
        
        'Third, Service Blueprint. This is unique to the SDD. A service blueprint maps the relationship between what the citizen sees and what happens behind the scenes. It has five layers: physical evidence, citizen actions, front-stage activities, back-stage activities and support systems. Neither the BRD nor the Gap Analysis provides this multi-layer view.',
        
        'Fourth, Service Channels and Touchpoints. The SDD defines which channels will be used (Telegram bot, web portal, ASAN Service centers) and what happens at each touchpoint. The BRD might mention channels, but the SDD designs the complete channel strategy with target volumes for each channel.',
        
        'Fifth, Multi-Agency Coordination Model. For government services, this is critical. The SDD defines which agencies are involved, what data they provide, what their role is and how they coordinate. The BRD mentions stakeholders but does not design the coordination model.',
        
        'Sixth, SLA and KPI Framework. The SDD defines specific service level agreements and key performance indicators with measurable targets. For example, SLA-03 says: "Application processing within 5 business days with 90% compliance." KPI-05 says: "Citizen satisfaction score of 4.0 out of 5.0." The BRD might mention success criteria, but the SDD creates a complete monitoring framework.',
        
        'Seventh, Service Portfolio Alignment. The SDD checks how the new service fits into the existing portfolio of government services and ensures alignment with national digital government standards. This strategic alignment is not covered by the BRD or Gap Analysis.',
        
        'So in summary: the BRD defines requirements, the Gap Analysis defines process improvements, and the SDD designs the complete citizen service experience including who uses it, how they experience it, what channels deliver it, how agencies coordinate, and how we measure success.',
    ]
)

# ─── Q45: What should an SDD contain? ───
add_qa_block(doc, 45,
    'What is the correct structure of a Service Design Document? What sections should it contain?',
    [
        'A Service Design Document follows a design thinking approach, not a traditional requirements approach. In my LMAS SDD, I used fifteen sections organized into four logical groups.',
        
        'The first group is Discovery. This includes sections 1 through 4. Section 1 is Executive Summary, which gives a one-page overview of the entire service design for busy decision-makers. Section 2 is Service Vision and Objectives, which defines the long-term vision, strategic objectives and measurable success metrics. Section 3 is Service Context and Stakeholder Analysis, which maps all stakeholders, their roles and how you communicate with each group. Section 4 is User Research and Personas, which documents your research findings, key insights and three to five citizen personas with their needs, behaviors and pain points.',
        
        'The second group is Design. This includes sections 5 through 9. Section 5 is As-Is Service Journey Analysis, which maps the current citizen journey stage by stage with pain points at each stage. Section 6 is To-Be Service Design, which describes the redesigned service journey with improvements at each stage and design principles that guide all decisions. Section 7 is Service Blueprint, which creates a five-layer view connecting citizen actions to front-stage, back-stage and support systems. Section 8 is Service Channels and Touchpoints, which defines the channel strategy and designs each interaction point. Section 9 is Multi-Agency Coordination Model, which defines how government agencies share data, their roles and the integration architecture.',
        
        'The third group is Standards. This includes sections 10 through 12. Section 10 is Service Standards, which defines SLA targets for each service stage and KPI indicators with specific measurable targets. Section 11 is Service Portfolio Alignment, which checks how the service fits into the broader government service portfolio and national digital standards. Section 12 is Process Architecture, which defines the three levels of process design: strategic, tactical and operational.',
        
        'The fourth group is Delivery. This includes sections 13 through 15. Section 13 is Implementation Roadmap, which plans the phased delivery approach with milestones and dependencies. Section 14 is Risks and Dependencies, which includes a risk register with mitigation strategies. Section 15 is Glossary, which defines all key terms used in the document.',
        
        'One important thing to note: the SDD is not a technical document. It does not contain database schemas, API endpoints or code-level specifications. Those belong in the SRS and API Specification. The SDD is a design document that focuses on the citizen experience and service delivery model.',
    ]
)

# ─── Q46: ToR vs SDD - what's the relationship? ───
add_qa_block(doc, 46,
    'You prepared both a ToR and an SDD for the LMAS project. How do these two documents relate to each other? Which one comes first and how do they connect?',
    [
        'The ToR comes first, and the SDD comes second. They have a parent-child relationship where the ToR sets the boundaries and the SDD fills in the design details.',
        
        'Think of it like building a house. The ToR is like the building permit and the project plan. It says: we are building a three-bedroom house on this plot of land, with this budget, this timeline and these people responsible. It defines what you are allowed to build and how the project will be managed.',
        
        'The SDD is like the architectural design. It says: the living room will be here, the kitchen will face east for morning sunlight, the entrance will have these features for wheelchair accessibility and the garden will be designed for family gatherings. It designs the experience of living in the house.',
        
        'In the LMAS project, the ToR says: "OBJ-01: Enable citizens to submit employment service applications through digital channels." This is a project objective with scope and timeline. The SDD then designs how this actually works: which channels, what the citizen journey looks like, what touchpoints the citizen interacts with, what happens behind the scenes and how success is measured.',
        
        'The ToR also defines the deliverables list. In my ToR, deliverable D-03 is the SDD itself. So the ToR creates the requirement to produce the SDD, defines who is responsible and sets the deadline. The SDD then fulfills this requirement.',
        
        'Another connection is the scope. The ToR defines what is in-scope and out-of-scope. The SDD must stay within these boundaries. For example, my ToR says that a mobile native application is out of scope for Phase 1. So the SDD designs the service for Telegram bot and web portal only, not for a mobile app. If the SDD started designing a mobile app experience, that would be a scope violation.',
        
        'The risk register in the ToR also connects to the SDD. For example, ToR risk R-03 says: "Low citizen adoption of digital channels." The SDD addresses this risk by designing a simplified onboarding experience in the citizen journey and by defining ASAN Service centers as a fallback channel for citizens who cannot use digital channels.',
        
        'So the relationship is: ToR defines the project frame and governance. SDD designs the service within that frame. ToR is the first document you write. SDD comes after, and it must respect the scope, timeline and constraints defined in the ToR.',
    ]
)

# ─── Q47: In what order do you prepare all BA documents? ───
add_qa_block(doc, 47,
    'You have prepared many BA documents: ToR, BRD, SDD, FRD, SRS, User Stories, API Specifications, BPMN models and UAT Plan. In what order do you prepare them and why?',
    [
        'There is a logical order, but in practice some documents overlap. Let me explain the ideal order and then how it works in reality.',
        
        'The ideal order follows a waterfall approach within each phase. Phase 1 starts with the ToR, because you cannot start a project without defining the scope, objectives, governance and timeline. The ToR gets approved by the steering committee. Then you write the BRD, which defines the business requirements in detail. Then you write the SDD, which designs the service experience. Then you create BPMN models for the As-Is and To-Be processes. All of these happen in Phase 1.',
        
        'Phase 2 is when you write the FRD, which converts business requirements into detailed functional requirements with user stories and acceptance criteria. Then you write the SRS, which adds technical system requirements. Then you write the API Specifications for all integrations. These happen in Phase 2.',
        
        'Phase 3 is when you prepare the UAT Plan and execute testing with stakeholders.',
        
        'So the ideal order is: ToR, BRD, SDD, BPMN, FRD, SRS, API Spec, UAT Plan.',
        
        'But in reality, some documents are written in parallel. For example, I usually start the BRD and the SDD at the same time because they inform each other. When I write a BRD requirement like "citizens must be able to track application status," this influences the SDD service journey design. And when I design the citizen journey in the SDD, I often discover new requirements that need to go into the BRD. So I work on both documents and update them as I go.',
        
        'The same is true for the FRD and SRS. The FRD defines the functional requirements, and the SRS adds the technical details. These two documents are closely connected, and I usually work on them together.',
        
        'One important rule I follow: the ToR must be completed and approved first before any other document. The ToR defines the project boundaries, and if you start writing the BRD without an approved ToR, you risk writing requirements for things that are out of scope or missing things that are in scope.',
        
        'Another rule: the UAT Plan is always the last document. You cannot plan testing until all requirements are defined and documented.',
    ]
)

# ─── Q48: How does the SDD relate to Service Design methodology? ───
add_qa_block(doc, 48,
    'The Innovation Agency uses Service Design methodology. How does your SDD document reflect Service Design methodology, and what Service Design tools and techniques did you use?',
    [
        'Service Design is a methodology that puts the citizen at the center of the design process. Instead of designing services from the institution perspective, you design them from the citizen perspective. My SDD uses several Service Design tools and techniques to apply this methodology.',
        
        'The first tool is User Research. In Service Design, you always start with research before design. In my SDD, section 4 documents the user research I conducted: 12 stakeholder interviews and analysis of 200+ citizen feedback records. This research identified six key findings about citizen pain points, which directly influenced all design decisions.',
        
        'The second tool is Personas. Personas are fictional characters that represent real user groups. They help the design team think about real people instead of abstract requirements. In my SDD, I created three personas based on the research data: a young job seeker, a mid-career worker and an employer. Each persona has specific needs, behaviors and pain points that the service design must address.',
        
        'The third tool is Service Journey Mapping. This is one of the most important Service Design tools. A journey map shows every step a citizen takes from the moment they become aware of a service to the moment they receive the result. In my SDD, sections 5 and 6 contain the As-Is and To-Be journey maps. The As-Is map shows seven stages of the current process with pain points at each stage. The To-Be map shows the redesigned seven stages with specific improvements.',
        
        'The fourth tool is Service Blueprinting. A service blueprint is like an x-ray of the service. It shows five layers: what the citizen sees (physical evidence), what the citizen does (actions), what happens in front of the citizen (front-stage), what happens behind the scenes (back-stage) and what systems support everything (support systems). In my SDD, section 7 contains the service blueprint that maps all five layers.',
        
        'The fifth tool is Life-Event Based Design. This is a Service Design approach where services are organized around key life events instead of government departments. Instead of saying "these are the services the employment agency provides," you say "when a citizen loses their job, these are all the services they need." In my SDD, I applied this approach by designing the service around the life event of unemployment rather than around agency departments.',
        
        'The sixth tool is Multi-Channel Design. Citizens should be able to access services through multiple channels: Telegram bot for quick interactions, web portal for complex tasks, ASAN centers for exceptional cases. My SDD section 8 designs the channel strategy with target volumes for each channel.',
        
        'The seventh tool is SLA and KPI Design. Service Design is not only about designing the service but also about measuring its performance. My SDD section 10 defines six SLAs and seven KPIs with specific measurable targets. This is the "how do we know the design works" part of Service Design.',
        
        'So my SDD is not just a document. It is the output of applying Service Design methodology to a real government service. Each section uses a specific Service Design tool to ensure the final service is citizen-centered, measurable and aligned across multiple agencies.',
    ]
)

# Save
output = '/home/z/my-project/ba-practice/Interview_Practice_BA_Portfolio_QA.docx'
doc.save(output)
print(f'Q&A document updated: {output}')
