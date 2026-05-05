#!/usr/bin/env python3
"""Generate Azerbaijani Cover Letter for Innovation Agency vacancy"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

def add_para(text, bold=False, size=11, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=Pt(6), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return p

# ── Header ──
add_para('ZAMİR JAMALOV', bold=True, size=16, space_after=Pt(2))
add_para('+994 55 207 7228  |  jamalov.zamir@gmail.com  |  Bakı, Azərbaycan',
         size=10, color=RGBColor(89, 89, 89), space_after=Pt(16))

# ── Date ──
add_para('6 may 2026-cı il', alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(18))

# ── Addressee ──
add_para('İnnovasiya və Rəqəmsal İnkişaf Agentliyi', space_after=Pt(2))
add_para('Bakı şəhəri', space_after=Pt(14))

# ── Subject ──
add_para(
    'Mövzu: "Dövlət xidmətlərinin dizaynı və təkmilləşdirilməsi üzrə '
    'Baş Mütəxəssis / Baş Biznes Analitik" vəzifəsinə namizədliyim haqqında',
    bold=True, space_after=Pt(14))

# ── Greeting ──
add_para('Hörmətli Rəhbərlik,', space_after=Pt(10))

# ── Body Paragraphs ──
body = [
    # P1: Introduction + motivation
    (
        'İnnovasiya və Rəqəmsal İnkişaf Agentliyində "Dövlət xidmətlərinin '
        'dizaynı və təkmilləşdirilməsi üzrə Baş Mütəxəssis" vəzifəsinə '
        'namizədliyim haqqında müraciət edirəm. İnformasiya texnologiyaları '
        'sahəsində 18 illik təcrübəm, o cümlədən dövlət xidmətlərinin '
        'digitallaşdırılması, çoxagentliyi inteqrasiya və biznes analitikası '
        'istiqamətindəki bilik və bacarıqlarım Agentliyin rəqəmsal dövlət '
        'xidmətlərini inkişaf etdirmək missiyasına mənaful töhfə verə '
        'biləcəyindən əminəm.'
    ),

    # P2: Government experience (GPP + DMA/EMAS)
    (
        'Karyeram ərzində Azərbaycanın dövlət xidmətləri ekosistemində '
        'birbaşa iştirak etmişəm. Azərbaycan Respublikasının Mərkəzi Bankında '
        'işləyərkən Dövlət Ödəniş Portalı (GPP) layihəsi çərçivəsində '
        '10-dan çox dövlət orqanının texniki inteqrasiyasını həyata '
        'keçirmişəm. Bu layihə mənə dövlət qurumlarının qarşılıqlı əlaqə '
        'mexanizmlərini, inteqrasiya prosesində yaranan sistemli çətinlikləri '
        'və strukturiləşdirilmiş analiz vasitəsilə həll yollarını dərindən '
        'başa düşmək imkanı vermişdir. Dövlət Məşğulluq Agentliyində (DMA) '
        'İnnovasiyalar şöbəsinin rəhbəri və EMAS (Məşğulluğun İdarəolunması '
        'Avtomatlaşdırma Sistemi) layihəsinin biznes analitiki kimi '
        'çalışmışam. 15 nəfərlik layihə komandası ilə birlikdə tələblər '
        'sənədləşməsini aparmış, sistemin ilkin inkişaf mərhələsində texniki '
        'komandalarla əlaqələndirmə funksiyasını yerinə yetirmişəm. '
        'Vətəndaşların dövlət xidmətlərinə çıxışını asanlaşdırmaq '
        'məqsədilə Telegram vasitəsilə real vaxt rejimində müraciət qəbulu '
        'və cavablandırma sistemi dizayn etmişəm. Həmçinin, idarəetmə '
        'şurası üçün veb-əsaslı monitoring paneli hazırlayaraq vətəndaş '
        'müraciətləri, xidmət çatdırılma müddətləri və keyfiyyət '
        'göstəricilərinin (KPI) şəffaf izlənilməsini təmin etmişəm.'
    ),

    # P3: Fintech experience + methodology
    (
        'Son illərdə fintech və e-kommersiya sahələrində fəaliyyət '
        'göstərərək beynəlxalq biznes analitikası standartlarını, xidmət '
        'dizaynı metodologiyalarını və müştəri səyahəti xəritələndirməsini '
        '(Customer Journey Mapping) praktiki olaraq tətbiq etmişəm. '
        'Tələblərin müəyyən edilməsi (BRD, FRD, SRS), As-Is / To-Be proses '
        'təhlili, xidmət çertyojunun (Service Blueprint) yaradılması, BPMN '
        'proses modelleməsi və tələblərin prioritetləşdirilməsi üzrə '
        '14-dən çox tam istehsalat səviyyəsində sənəd hazırlamışam. Bu '
        'təcrübə metodologiya mənim üçün yalnız nəzəri bilik deyil, hər gün '
        'tətbiq etdiyim praktikadır.'
    ),

    # P4: Why this agency + life-event service design
    (
        'Agentliyin fəaliyyət istiqaməti — həyat hadisələrinə əsaslanan '
        'xidmət dizaynı, ucdan-uca proses arxitekturası və çoxagentliyi '
        'əlaqələndirmə — mənim peşəkar maraqlarımla tam uyğun gəlir. Bu '
        'konsepsiyalar mənim üçün abstrakt deyil: GPP çoxagentliyi '
        'inteqrasiyasında, EMAS məşğulluq xidmətinin digitallaşdırılmasında '
        'və vətəndaşlara yönəlmiş rəqəmsal xidmət kanallarının yaradılmasında '
        'bu prinsipləri reallıqda tətbiq etmişəm. Dövlət qurumlarının işləmə '
        'qaydalarını, müxtəlif agentliklərin maraqlı sahibləri ilə necə '
        'əlaqə saxlamaq lazım olduğunu və vətəndaşı mərkəzə qoyan xidmətləri '
        'necə dizayn etmək olarını bilirəm.'
    ),

    # P5: Closing
    (
        'Biznes analitikası metodologiyası, dövlət sektoru təcrübəsi və '
        'texniki arxaplanımın birləşməsi Agentliyin rəqəmsal transformasiya '
        'təşəbbüslərinə birinci gündən töhfə verməyə imkan verəcəyinə '
        'inanıram. Müsahibə imkanı üçün təşəkkür edirəm.'
    ),
]

for text in body:
    add_para(text, size=11, space_after=Pt(8))

# ── Closing ──
add_para('Hörmətlə,', space_before=Pt(14), space_after=Pt(4))
add_para('Zamir Jamalov', bold=True, size=12, space_after=Pt(0))

# Save
output = '/home/z/my-project/ba-practice/Zamir_Jamalov_Cover_Letter_Innovation_Agency_AZ.docx'
doc.save(output)
print(f'Cover Letter saved: {output}')
