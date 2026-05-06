#!/usr/bin/env python3
"""Generate Innovation Agency Cover Letter v2 - English B1 level"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

DARK = RGBColor(33, 37, 41)
GRAY = RGBColor(89, 89, 89)

def add_para(text, bold=False, size=11, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=Pt(6), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return p

# Header
add_para('ZAMIR JAMALOV', bold=True, size=16, space_after=Pt(2))
add_para('+994 55 207 7228  |  jamalov.zamir@gmail.com  |  Baku, Azerbaijan',
         size=10, color=GRAY, space_after=Pt(16))

# Date
add_para('May 6, 2026', alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(18))

# Addressee
add_para('Innovation and Digital Development Agency', space_after=Pt(2))
add_para('Baku, Azerbaijan', space_after=Pt(14))

# Subject
add_para(
    'Subject: Application for Senior Business Analyst / Lead Specialist - '
    'Design and Improvement of Public Services',
    bold=True, space_after=Pt(14))

# Greeting
add_para('Dear Hiring Manager,', space_after=Pt(10))

# Body paragraphs - B1 English level
body = [
    (
        'I am writing to express my interest in the Senior Business Analyst / '
        'Lead Specialist position at the Innovation and Digital Development Agency. '
        'With 18 years in IT, including direct experience in government service '
        'digitization, multi-agency coordination and business analysis, I believe '
        'my background is well suited for the Agency\'s mission to improve public '
        'services through digital transformation.'
    ),

    (
        'My career has been built around a consistent theme: connecting government '
        'institutions with technology. At the Central Bank of Azerbaijan, I '
        'coordinated the integration of 10+ government organizations into the '
        'Government Payment Portal (GPP). This required defining data exchange '
        'requirements for each agency, building cross-system middleware, and working '
        'closely with multiple stakeholders to ensure seamless payment processing at '
        'national scale. This experience gave me a deep understanding of how government '
        'systems interact and where integration challenges arise.'
    ),

    (
        'At the State Employment Agency, I led the Innovation Department and served '
        'as Business Analyst for the Labour and Employment Subsystem (LMAS). I worked '
        'with a 15-member project team to digitize the employment service process. '
        'Specifically, I analyzed the full citizen service journey from application to '
        'result, identified pain points in the existing process, and designed the '
        'end-to-end architecture for the new system. Beyond traditional BA work, I '
        'designed a Telegram-based citizen service channel that enabled real-time '
        'application submission, directly improving service accessibility for citizens '
        'who could not visit the office in person. I also built a real-time monitoring '
        'dashboard for the management board, which provided transparent tracking of '
        'citizen applications, response times and service quality indicators.'
    ),

    (
        'In my current role at Embafinans, I lead business analysis for fintech '
        'products. I author BRD, FRD and SRS documents, write User Stories with '
        'Gherkin Acceptance Criteria, and define REST API specifications. I have '
        'hands-on experience with As-Is / To-Be process analysis, BPMN modeling, '
        'and backlog prioritization using the RICE framework. These methodologies '
        'are directly applicable to the Agency\'s work of analyzing and redesigning '
        'public services.'
    ),

    (
        'What attracts me most to this position is the Agency\'s focus on end-to-end '
        'process architecture, multi-agency coordination and citizen-centered service '
        'design. These are not abstract concepts for me. I have applied them through '
        'the GPP multi-agency integration, the LMAS employment service digitization, '
        'and the citizen-facing Telegram channel. I understand how government '
        'institutions operate, how to engage stakeholders across different agencies, '
        'and how to design services that put citizens at the center.'
    ),

    (
        'I am confident that my combination of business analysis methodology, '
        'government sector experience and technical background would allow me to '
        'contribute meaningfully to the Agency\'s digital transformation initiatives '
        'from day one. I would welcome the opportunity to discuss how my experience '
        'can support the Agency\'s goals.'
    ),
]

for text in body:
    add_para(text, size=11, space_after=Pt(8))

# Closing
add_para('Sincerely,', space_before=Pt(14), space_after=Pt(4))
add_para('Zamir Jamalov', bold=True, size=12)

output = '/home/z/my-project/ba-practice/Zamir_Jamalov_Cover_Letter_Innovation_Agency_EN.docx'
doc.save(output)
print(f'Cover Letter EN saved: {output}')
