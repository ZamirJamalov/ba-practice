#!/usr/bin/env python3
"""Generate updated Innovation Agency CV v2 with Service Design + weaknesses covered"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# ── Helpers ──
DARK = RGBColor(33, 37, 41)
ACCENT = RGBColor(0, 90, 156)       # professional blue
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = RGBColor(120, 120, 120)

def add_line(text, bold=False, size=9.5, color=DARK, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=Pt(2), space_before=Pt(0), font_name='Calibri'):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = Pt(12.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run.font.color.rgb = color
    return p

def add_bullet(text, size=9, color=DARK, indent=0.25):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    run = p.add_run('•  ' + text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = color
    return p

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    # thin horizontal line via bottom border on empty paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'AAAAAA'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

# ════════════════════════════════════════════════════════════════════
# HEADER
# ════════════════════════════════════════════════════════════════════
add_line('ZAMİR JAMALOV', bold=True, size=15, color=DARK, space_after=Pt(2))
add_line('Business Analyst  |  18 Years in IT', size=9.5, color=ACCENT, space_after=Pt(2))
add_line('+994 55 207 7228  |  jamalov.zamir@gmail.com  |  Bakı, Azərbaycan', size=9, color=GRAY, space_after=Pt(4))

add_divider()

# ════════════════════════════════════════════════════════════════════
# PROFILE SUMMARY
# ════════════════════════════════════════════════════════════════════
add_line('PROFIL XÜLASƏSİ', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))
add_line(
    'Biznes analitikası, sistem inteqrasiyası və dövlət sektoru təcrübəsini '
    'birləşdirən 18 illik IT peşəkar. 14-dən çox istehsalat səviyyəsində BA sənədinin '
    'müəllifi: BRD, FRD, SRS, User Stories, API spesifikasiyaları, As-Is / To-Be təhlil. '
    'Dövlət Ödəniş Portalı (GPP) çərçivəsində 10+ qurumun inteqrasiyası, EMAS məşğulluq '
    'sisteminin digitallaşdırılması, vətəndaşlara yönəlmiş rəqəmsal xidmət kanallarının '
    'yaradılması. Xidmət dizaynı (Customer Journey Mapping, Service Blueprint) və ucdan-uca '
    'proses arxitekturası üzrə praktiki təcrübə.',
    size=8.5, color=DARK, space_after=Pt(4))

add_divider()

# ════════════════════════════════════════════════════════════════════
# CORE SKILLS
# ════════════════════════════════════════════════════════════════════
add_line('ƏSAS BACARIQLAR', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))

add_line('Xidmət Dizaynı:', bold=True, size=8.5, color=DARK, space_after=Pt(1))
add_line(
    'Customer Journey Mapping  |  Service Blueprint  |  Design Thinking  |  '
    'As-Is / To-Be Xidmət Analizi  |  Vətəndaşa Yönəlmiş Xidmət Dizaynı  |  '
    'Həyat Hadisələrinə Əsaslanan Xidmət Modelləşdirməsi  |  SLA / KPI Monitorinqi',
    size=8, color=DARK, space_after=Pt(3))

add_line('Biznes Analitikası:', bold=True, size=8.5, color=DARK, space_after=Pt(1))
add_line(
    'BRD / FRD / SRS  |  User Stories & Acceptance Criteria  |  '
    'BPMN (As-Is / To-Be)  |  UML & Sequence Diagrams  |  Gap Analysis  |  '
    'Stakeholder İntervyuları  |  Backlog Prioritization (RICE)  |  UAT',
    size=8, color=DARK, space_after=Pt(3))

add_line('Texniki:', bold=True, size=8.5, color=DARK, space_after=Pt(1))
add_line(
    'REST API & JSON  |  Swagger / OpenAPI 3.0  |  SQL  |  '
    'SDLC  |  Sistem İnteqrasiyası  |  Data-Driven Analysis',
    size=8, color=DARK, space_after=Pt(3))

add_line('Proses & Alətlər:', bold=True, size=8.5, color=DARK, space_after=Pt(1))
add_line(
    'Agile / Scrum  |  Jira  |  Confluence  |  Çoxfunksional Əlaqələndirmə  |  '
    'Dövlət Sektoru Stakeholder İdarəetməsi  |  Proses Digitallaşdırılması',
    size=8, color=DARK, space_after=Pt(4))

add_divider()

# ════════════════════════════════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ════════════════════════════════════════════════════════════════════
add_line('PEŞƏKAR TƏCRÜBƏ', bold=True, size=9.5, color=ACCENT, space_after=Pt(4))

# ── 1. Embafinans ──
add_line('Embafinans  |  Baş IT Biznes Analitiki', bold=True, size=9, color=DARK, space_after=Pt(1))
add_line('2025 – Hazırkı', size=8, color=LIGHT_GRAY, space_after=Pt(2))
add_bullet('Fintech məhsulları üzrə biznes analitikasının rəhbərliyi: BNPL credit scoring, ödəniş gateway inteqrasiyası, kredit izləmə dashboardu')
add_bullet('BRD, FRD, SRS sənədlərinin yaradılması; User Stories (Gherkin AC) ilə sprintlər boyu traceability-in saxlanması')
add_bullet('REST API spesifikasiyalarının Swagger / OpenAPI 3.0-da müəyyən edilməsi, sequence diaqramları ilə inteqrasiya axınlarının modelleməsi')
add_bullet('UAT icrasının biznes sahibləri ilə əlaqələndirilməsi, bug triage iclaslarının rəhbərliyi, release dövrlərində zamanında sign-off')

# ── 2. Kapital Bank / Birbonus ──
add_line('Kapital Bank / Birbonus  |  IT Biznes Analitiki', bold=True, size=9, color=DARK, space_before=Pt(4), space_after=Pt(1))
add_line('2024 – 2025', size=8, color=LIGHT_GRAY, space_after=Pt(2))
add_bullet('Müştəri sadiqlik bonus sistemi (Birmarket) dizaynı: alış-veriş zamanı xal qazanma və partner satıcılarda xərcləmə modeli')
add_bullet('Stakeholder sessiyaları ilə qazanma qaydaları, uyğunluq meyarları və partner hesablaşma workflow-larının müəyyən edilməsi')

# ── 3. Umico ──
add_line('Umico  |  İnteqrasiya Mütəxəssisi', bold=True, size=9, color=DARK, space_before=Pt(4), space_after=Pt(1))
add_line('2022 – 2024', size=8, color=LIGHT_GRAY, space_after=Pt(2))
add_bullet('25+ partner şirkətin Umico ekosistemə API inteqrasiyası: inteqrasiya spesifikasiyalarının müəyyən edilməsi və texniki icrasının əlaqələndirilməsi')
add_bullet('Backend xüsusiyyətlərin inkişaf etdirilməsi (PostgreSQL), L2 istehsalat insidentlərinin həlli, partner komandalarına dəstək')

# ── 4. DMA ──
add_line('Dövlət Məşğulluq Agentliyi  |  İnnovasiyalar Şöbəsinin Rəhbəri & Biznes Analitik', bold=True, size=9, color=DARK, space_before=Pt(4), space_after=Pt(1))
add_line('2021 – 2022', size=8, color=LIGHT_GRAY, space_after=Pt(2))
add_bullet('İnnovasiyalar şöbəsinin rəhbərliyi (2 nəfər), 15 nəfərlik EMAS layihə komandası ilə əlaqələndirmə, metodoloji bələdçilik və bilik transferi')
add_bullet('EMAS (Məşğulluğun İdarəolunması Avtomatlaşdırma Sistemi) üzrə biznes analitik: tələblərin toplanması, sənədləşməsi, texniki komandaya təhvil verilməsi')
add_bullet('Vətəndaşların xidmətə çıxış yollarının dizaynı: Telegram vasitəsilə real vaxt rejimində müraciət qəbulu və cavablandırma sistemi — xidmətə çıxışın asanlaşdırılması')
add_bullet('İdarəetmə şurası üçün real vaxt monitoring dashboardu: vətəndaş müraciətləri, cavab müddətləri, xidmət keyfiyyəti KPI-larının şəffaf izlənməsi')
add_bullet('Vətəndaşın məşğulluq xidmətinə müraciətindən nəticə əldə etməyədək olan tam xidmət yolunun (service journey) təhlili və təkmilləşdirilməsi')

# ── 5. Central Bank ──
add_line('Azərbaycan Respublikasının Mərkəzi Bankı  |  İnteqrasiya Developer', bold=True, size=9, color=DARK, space_before=Pt(4), space_after=Pt(1))
add_line('2007 – 2012', size=8, color=LIGHT_GRAY, space_after=Pt(2))
add_bullet('Dövlət Ödəniş Portalı (GPP) layihəsi: 10+ dövlət qurumunun portal-a texniki inteqrasiyası, məlumat mübadiləsi spesifikasiyalarının müəyyən edilməsi')
add_bullet('Dövlət qurumları arasında cross-system middleware-in inkişaf etdirilməsi, milli miqyasda avtomatlaşdırılmış ödəniş emalının təmin edilməsi')

add_divider()

# ════════════════════════════════════════════════════════════════════
# PROFESSIONAL DEVELOPMENT
# ════════════════════════════════════════════════════════════════════
add_line('PEŞƏKAR İNKİŞAF', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))
add_line(
    'Biznes analitikası metodologiyasını real layihələrdə praktiki tətbiq etməklə mənimsəmişəm: '
    '14-dən çox istehsalat səviyyəsində BA sənədi (BRD, FRD, SRS, User Stories, API Specs, BPMN, UAT) '
    'fintech, e-kommersiya və dövlət xidmətləri sahələrində. Biznes proses və xidmət dizaynı '
    'metodologiyalarının hər ikisində praktik təcrübəyə malikəm.',
    size=8.5, color=DARK, space_after=Pt(4))

add_divider()

# ════════════════════════════════════════════════════════════════════
# ADDITIONAL BACKGROUND
# ════════════════════════════════════════════════════════════════════
add_line('ƏLAVƏ PEŞƏKAR ARXAPLAN', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))
add_line(
    '10+ illik proqram mühəndisliyi və data analitikası təcrübəsi Azərbaycanın bank və dövlət '
    'sektorlarında: core banking inkişafı (Zaminbank, Unibank, Bank of Baku, Rabita Bank), backend '
    'development (Umico), və dövlət xidmətləri (ASAN Service). Bu təməl enterprise sistemlər, '
    'verilənlər bazası arxitekturaları (Oracle, MSSQL, PostgreSQL, MongoDB), dövlət xidmət '
    'çatdırılma nümunələri və çoxtəşkilatlı inteqrasiya barədə dərin anlayış təmin edir.',
    size=8.5, color=DARK, space_after=Pt(4))

add_divider()

# ════════════════════════════════════════════════════════════════════
# EDUCATION
# ════════════════════════════════════════════════════════════════════
add_line('TƏHSİL', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))
add_line('Bakı Dövlət Universiteti  —  Tətbiqi Riyaziyyat üzrə Bakalavr', size=9, color=DARK, space_after=Pt(4))

add_line('DİLLƏR', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))
add_line('Azərbaycan (Ana dili)  |  Rus (Sərbəst)  |  İngilis (Peşəkar / Texniki Sənədləşmə)', size=9, color=DARK)

# ── Save ──
output = '/home/z/my-project/ba-practice/Zamir_Jamalov_CV_BA_Innovation_Agency.docx'
doc.save(output)
print(f'CV v2 saved: {output}')
