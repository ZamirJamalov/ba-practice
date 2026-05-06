#!/usr/bin/env python3
"""Generate Innovation Agency CV v3 - English B1 with corrected LMAS and Applied Service Design strategy"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

DARK = RGBColor(33, 37, 41)
ACCENT = RGBColor(0, 90, 156)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = RGBColor(120, 120, 120)

def add_line(text, bold=False, size=9.5, color=DARK, alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=Pt(2), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = Pt(13)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = color
    return p

def add_bullet(text, size=9, color=DARK, indent=0.25):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(12.5)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = color
    return p

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): 'AAAAAA'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

# HEADER
add_line('ZAMIR JAMALOV', bold=True, size=15, color=DARK, space_after=Pt(2))
add_line('Business Analyst  |  Government Service Design & Multi-Agency Coordination',
         size=9.5, color=ACCENT, space_after=Pt(2))
add_line('+994 55 207 7228  |  jamalov.zamir@gmail.com  |  Baku, Azerbaijan',
         size=9, color=GRAY, space_after=Pt(4))
add_divider()

# PROFILE SUMMARY
add_line('PROFILE SUMMARY', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))
add_line(
    'Business Analyst with 18 years in IT, specializing in government service '
    'digitization, multi-agency coordination, and end-to-end process design. '
    'Practical experience in service journey analysis, citizen-centered digital '
    'service design, and requirements documentation (BRD, FRD, SRS, API '
    'Specifications). Strong background in aligning cross-functional teams '
    'across government organizations to deliver digital public services.',
    size=8.5, color=DARK, space_after=Pt(4))
add_divider()

# CORE SKILLS
add_line('CORE SKILLS', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))

add_line('Service Design & Analysis:', bold=True, size=8.5, color=DARK, space_after=Pt(1))
add_line(
    'Applied Service Design  |  Service Journey Analysis  |  Citizen-Centered '
    'Service Design  |  Multi-Agency Service Coordination  |  As-Is / To-Be '
    'Service Analysis  |  SLA / KPI Monitoring  |  Process Digitization',
    size=8, color=DARK, space_after=Pt(3))

add_line('Business Analysis:', bold=True, size=8.5, color=DARK, space_after=Pt(1))
add_line(
    'BRD / FRD / SRS  |  User Stories & Acceptance Criteria  |  BPMN (As-Is / '
    'To-Be)  |  UML & Sequence Diagrams  |  Gap Analysis  |  Stakeholder '
    'Interviews  |  Backlog Prioritization (RICE)  |  UAT',
    size=8, color=DARK, space_after=Pt(3))

add_line('Technical:', bold=True, size=8.5, color=DARK, space_after=Pt(1))
add_line(
    'REST API & JSON  |  Swagger / OpenAPI 3.0  |  SQL  |  SDLC  |  System '
    'Integration  |  Data-Driven Analysis',
    size=8, color=DARK, space_after=Pt(3))

add_line('Process & Tools:', bold=True, size=8.5, color=DARK, space_after=Pt(1))
add_line(
    'Agile / Scrum  |  Jira  |  Confluence  |  Cross-Functional Coordination  |  '
    'Government Stakeholder Management',
    size=8, color=DARK, space_after=Pt(4))
add_divider()

# PROFESSIONAL EXPERIENCE - 5 entries + Additional Background
add_line('PROFESSIONAL EXPERIENCE', bold=True, size=9.5, color=ACCENT, space_after=Pt(4))

# 1. Embafinans (most recent)
add_line('Embafinans  |  Lead IT Business Analyst',
         bold=True, size=9, color=DARK, space_after=Pt(1))
add_line('2025 - Present', size=8, color=LIGHT_GRAY, space_after=Pt(2))
add_bullet(
    'Leading business analysis for fintech products including BNPL credit scoring, '
    'payment gateway integration, and loan tracking dashboard')
add_bullet(
    'Authoring BRD, FRD and SRS documents; writing User Stories with Gherkin '
    'Acceptance Criteria and maintaining traceability across sprints')
add_bullet(
    'Defining REST API specifications in Swagger / OpenAPI 3.0 and coordinating '
    'UAT execution with business stakeholders')

# 2. Kapital Bank
add_line('Kapital Bank / Birbonus  |  IT Business Analyst',
         bold=True, size=9, color=DARK, space_before=Pt(4), space_after=Pt(1))
add_line('2024 - 2025', size=8, color=LIGHT_GRAY, space_after=Pt(2))
add_bullet(
    'Designed a customer loyalty bonus system enabling shoppers to earn rewards '
    'on purchases and redeem across partner merchants')
add_bullet(
    'Conducted stakeholder sessions to define earning rules, eligibility criteria, '
    'and partner settlement workflows')

# 3. Umico - Integration
add_line('Umico  |  Integration Specialist',
         bold=True, size=9, color=DARK, space_before=Pt(4), space_after=Pt(1))
add_line('2022 - 2024', size=8, color=LIGHT_GRAY, space_after=Pt(2))
add_bullet(
    'Managed API integration of 25+ partner companies into the Umico ecosystem, '
    'defining integration specifications and coordinating technical implementation')
add_bullet(
    'Developed backend features using PostgreSQL, resolved L2 production incidents, '
    'and supported partner development teams')

# 4. DMA - Section Lead (most relevant for this position)
add_line('State Employment Agency  |  Innovation Section Lead & Business Analyst',
         bold=True, size=9, color=DARK, space_before=Pt(4), space_after=Pt(1))
add_line('2021 - 2022', size=8, color=LIGHT_GRAY, space_after=Pt(2))
add_bullet(
    'Led the digitization of the Labour and Employment Subsystem (LMAS): '
    'analyzed the full citizen life-event service journey from application to '
    'result, defined the methodology and end-to-end process architecture, and '
    'authored requirements documentation')
add_bullet(
    'Designed a citizen service channel using Telegram bot for real-time '
    'application submission, reducing physical office visits and improving accessibility')
add_bullet(
    'Built a real-time monitoring dashboard tracking citizen applications, '
    'response times (SLA), and service quality indicators (KPI)')
add_bullet(
    'Coordinated a 15-member cross-functional project team and managed '
    'multi-stakeholder communication across government departments')

# 5. Central Bank
add_line('Central Bank of Azerbaijan  |  Integration Specialist',
         bold=True, size=9, color=DARK, space_before=Pt(4), space_after=Pt(1))
add_line('2007 - 2012', size=8, color=LIGHT_GRAY, space_after=Pt(2))
add_bullet(
    'Coordinated the integration of 10+ government organizations into the '
    'single-window Government Payment Portal (GPP): defined data exchange '
    'requirements and integration specifications for each agency')
add_bullet(
    'Designed cross-system data exchange middleware between government institutions, '
    'enabling automated payment processing at national scale')

add_divider()

# PROFESSIONAL DEVELOPMENT
add_line('PROFESSIONAL DEVELOPMENT', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))
add_line(
    'Acquired business analysis methodology through practical application across '
    'fintech, e-commerce and government digitization projects. Produced BA '
    'documentation including BRD, FRD, SRS, User Stories, API Specifications, '
    'BPMN process models and UAT plans in real production environments.',
    size=8.5, color=DARK, space_after=Pt(4))
add_divider()

# ADDITIONAL PROFESSIONAL BACKGROUND
add_line('ADDITIONAL PROFESSIONAL BACKGROUND', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))
add_line(
    'Enterprise system design and implementation experience (2013-2021) across '
    "Azerbaijan's banking and public sectors: core banking system architecture and "
    'support at Zaminbank; mobile banking system design and development (UMobileBank) '
    'at Unibank; banking system modernization at Bank of Baku and Rabita Bank; '
    'e-commerce platform backend at Umico; and public service data analytics at '
    'ASAN Service. This period built deep understanding of how complex enterprise '
    'systems are designed, integrated and maintained at scale.',
    size=8.5, color=DARK, space_after=Pt(4))
add_divider()

# EDUCATION & LANGUAGES
add_line('EDUCATION', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))
add_line('Baku State University  -  Bachelor of Science in Applied Mathematics',
         size=9, color=DARK, space_after=Pt(4))

add_line('LANGUAGES', bold=True, size=9.5, color=ACCENT, space_after=Pt(2))
add_line('Azerbaijani (Native)  |  Russian (Fluent)  |  English (Professional Working Proficiency)',
         size=9, color=DARK)

output = '/home/z/my-project/ba-practice/Zamir_Jamalov_CV_Innovation_Agency_EN.docx'
doc.save(output)
print(f'CV EN saved: {output}')
