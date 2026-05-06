#!/usr/bin/env python3
"""Generate Service Design Document (SDD) for LMAS - English B1 level"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# Colors
ACCENT = RGBColor(0, 90, 156)
DARK = RGBColor(33, 37, 41)
GRAY = RGBColor(89, 89, 89)
WHITE = RGBColor(255, 255, 255)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = ACCENT
        run.font.name = 'Calibri'
    return h

def add_heading_2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'
    return h

def add_heading_3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'
    return h

def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    return p

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, (text, width_pct) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if header:
            run.bold = True
            run.font.color.rgb = WHITE
            set_cell_shading(cell, '005A9C')
        else:
            run.font.color.rgb = DARK
    return row

def add_simple_table(headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = table.add_row()
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, '005A9C')
    # Data
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK
    return table


# ======================================================================
# COVER PAGE
# ======================================================================
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run('STATE EMPLOYMENT AGENCY')
run.font.size = Pt(20)
run.font.color.rgb = ACCENT
run.font.name = 'Calibri'
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(20)
run = subtitle.add_run('Service Design Document (SDD)')
run.font.size = Pt(28)
run.font.color.rgb = DARK
run.font.name = 'Calibri'
run.bold = True

project = doc.add_paragraph()
project.alignment = WD_ALIGN_PARAGRAPH.CENTER
project.paragraph_format.space_after = Pt(30)
run = project.add_run('Labour and Employment Subsystem (LMAS)\nEnd-to-End Service Design')
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

# Meta info table
meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Version', '1.0'),
    ('Date', 'May 2026'),
    ('Author', 'Business Analysis Team'),
    ('Classification', 'Internal'),
    ('Status', 'Draft'),
]
for i, (label, value) in enumerate(meta_data):
    meta_table.rows[i].cells[0].text = label
    meta_table.rows[i].cells[1].text = value
    for j in range(2):
        for p in meta_table.rows[i].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                run.font.color.rgb = DARK
            if j == 0:
                p.runs[0].bold = True

doc.add_page_break()

# ======================================================================
# TABLE OF CONTENTS (placeholder)
# ======================================================================
add_heading_1('Table of Contents')
toc_items = [
    '1. Executive Summary',
    '2. Service Vision and Objectives',
    '3. Service Context and Stakeholder Analysis',
    '4. User Research and Personas',
    '5. As-Is Service Journey Analysis',
    '6. To-Be Service Design',
    '7. Service Blueprint',
    '8. Service Channels and Touchpoints',
    '9. Multi-Agency Coordination Model',
    '10. Service Standards (SLA and KPI)',
    '11. Service Portfolio Alignment',
    '12. Process Architecture',
    '13. Implementation Roadmap',
    '14. Risks and Dependencies',
    '15. Appendix: Glossary',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK

doc.add_page_break()

# ======================================================================
# 1. EXECUTIVE SUMMARY
# ======================================================================
add_heading_1('1. Executive Summary')

add_body(
    'This Service Design Document describes the end-to-end design of the '
    'Labour and Employment Subsystem (LMAS), a digital government service that '
    'allows citizens to access employment-related services through online '
    'channels. The service is being designed for the State Employment Agency '
    'of Azerbaijan with the goal of improving service accessibility, reducing '
    'processing time and increasing citizen satisfaction.'
)

add_body(
    'The current employment service process requires citizens to visit '
    'government offices in person, submit paper documents and wait for manual '
    'processing. This approach creates bottlenecks, limits accessibility for '
    'citizens outside of Baku and does not provide real-time status tracking. '
    'The LMAS project aims to digitize the full citizen service journey from '
    'application submission to final result delivery.'
)

add_body(
    'The service design follows a citizen-centered approach, using '
    'life-event-based service grouping where employment services are organized '
    'around key life events such as job search, unemployment registration and '
    'labor contract management. The design includes multiple digital channels '
    '(web portal, Telegram bot), a real-time monitoring dashboard and '
    'integration with other government systems through a single-window '
    'architecture.'
)

add_heading_3('Key Design Decisions')

key_decisions = [
    ('Citizen journey model', 'Life-event-based service journey from application to result'),
    ('Primary digital channel', 'Telegram bot for application submission and status tracking'),
    ('Secondary channel', 'Web portal for detailed service access and document upload'),
    ('Monitoring', 'Real-time management dashboard with SLA and KPI tracking'),
    ('Integration', 'Multi-agency data exchange with 5+ government organizations'),
    ('Methodology', 'End-to-end process architecture with As-Is / To-Be analysis'),
]
add_simple_table(['Decision Area', 'Selected Approach'], key_decisions)

doc.add_paragraph()

# ======================================================================
# 2. SERVICE VISION AND OBJECTIVES
# ======================================================================
add_heading_1('2. Service Vision and Objectives')

add_heading_2('2.1 Service Vision')
add_body(
    'To provide every citizen in Azerbaijan with accessible, transparent and '
    'efficient employment services through digital channels, reducing the need '
    'for physical office visits while maintaining high service quality standards.'
)

add_heading_2('2.2 Strategic Objectives')

objectives = [
    ('OBJ-01', 'Increase digital service adoption rate to 60% within the first year of operation', 'High'),
    ('OBJ-02', 'Reduce average service processing time from 14 days to 5 business days', 'High'),
    ('OBJ-03', 'Eliminate the requirement for physical office visits for standard employment services', 'Medium'),
    ('OBJ-04', 'Provide real-time application status tracking for 100% of submitted applications', 'High'),
    ('OBJ-05', 'Achieve citizen satisfaction score of 4.0 or higher (out of 5.0)', 'Medium'),
    ('OBJ-06', 'Establish SLA compliance rate of 90% or above for all service stages', 'High'),
]
add_simple_table(['ID', 'Objective', 'Priority'], objectives)

doc.add_paragraph()

add_heading_2('2.3 Success Metrics')

metrics = [
    ('Digital adoption rate', 'Percentage of applications submitted through digital channels', '60%', '12 months'),
    ('Average processing time', 'Calendar days from application to final decision', '5 days', '6 months'),
    ('Citizen satisfaction score', 'Average rating from post-service survey', '4.0 / 5.0', '12 months'),
    ('SLA compliance rate', 'Percentage of applications processed within SLA target', '90%', '6 months'),
    ('Office visit reduction', 'Percentage decrease in physical office visits', '70%', '12 months'),
    ('System uptime', 'Percentage of time digital channels are available', '99.5%', 'Ongoing'),
]
add_simple_table(['Metric', 'Description', 'Target', 'Timeline'], metrics)

doc.add_paragraph()

# ======================================================================
# 3. SERVICE CONTEXT AND STAKEHOLDER ANALYSIS
# ======================================================================
add_heading_1('3. Service Context and Stakeholder Analysis')

add_heading_2('3.1 Organizational Context')
add_body(
    'The State Employment Agency is responsible for managing employment '
    'services across Azerbaijan. The agency operates through regional offices '
    'where citizens can access services in person. The LMAS project is part '
    'of the broader digital government transformation strategy led by the '
    'Innovation and Digital Development Agency.'
)

add_body(
    'The service ecosystem includes multiple government organizations that '
    'contribute data and participate in the employment service process. '
    'These include the Ministry of Labour and Social Protection, the State '
    'Social Protection Fund, the Ministry of Education and local municipality '
    'offices. Coordinating across these organizations is a critical success '
    'factor for the LMAS project.'
)

add_heading_2('3.2 Stakeholder Map')

stakeholders = [
    ('Citizens (Job seekers)', 'Primary user', 'Submit applications, track status, receive results', 'High'),
    ('Employers', 'Secondary user', 'Post vacancies, review candidate profiles', 'Medium'),
    ('State Employment Agency', 'Service owner', 'Process applications, manage service delivery', 'High'),
    ('Ministry of Labour', 'Policy body', 'Set employment policies, regulatory oversight', 'High'),
    ('Social Protection Fund', 'Data provider', 'Provide social insurance and benefit data', 'Medium'),
    ('Innovation Agency', 'Strategic partner', 'Digital transformation methodology, standards', 'Medium'),
    ('IT Department', 'Technical team', 'System development, infrastructure, maintenance', 'High'),
    ('Regional Offices', 'Service delivery', 'Handle cases requiring physical presence', 'Medium'),
]
add_simple_table(
    ['Stakeholder', 'Role', 'Interaction', 'Influence'],
    stakeholders
)

doc.add_paragraph()

add_heading_2('3.3 Stakeholder Communication Plan')

comm = [
    ('State Employment Agency management', 'Bi-weekly steering committee', 'Project status, risks, decisions'),
    ('Ministry of Labour', 'Monthly alignment meeting', 'Policy updates, regulatory changes'),
    ('Innovation Agency', 'Monthly progress review', 'Methodology alignment, standards compliance'),
    ('IT Department', 'Weekly sprint review', 'Technical progress, blockers'),
    ('Regional Offices', 'Monthly feedback session', 'User feedback, process improvement'),
]
add_simple_table(['Stakeholder', 'Frequency', 'Topics'], comm)

doc.add_paragraph()

# ======================================================================
# 4. USER RESEARCH AND PERSONAS
# ======================================================================
add_heading_1('4. User Research and Personas')

add_heading_2('4.1 Research Approach')
add_body(
    'User research was conducted through stakeholder interviews with 12 '
    'employees of the State Employment Agency and analysis of 200+ citizen '
    'feedback records from the existing service delivery process. The research '
    'identified key pain points, user needs and behavioral patterns that '
    'directly influenced the service design decisions.'
)

add_heading_2('4.2 Key Research Findings')

findings = [
    ('Accessibility barrier', '65% of citizens outside Baku find it difficult to visit regional offices due to distance and travel costs'),
    ('Lack of transparency', '78% of applicants do not know the status of their application after submission'),
    ('Document complexity', '52% of applications are rejected due to incomplete or incorrect documentation'),
    ('Processing delays', 'Average processing time is 14 calendar days, with some cases taking over 30 days'),
    ('Limited channels', 'Currently only in-person service delivery is available, no digital alternatives'),
    ('Repeat visits', '40% of citizens visit the office more than once for a single service request'),
]
add_simple_table(['Finding', 'Detail'], findings)

doc.add_paragraph()

add_heading_2('4.3 Citizen Personas')

add_heading_3('Persona 1: Young Job Seeker')
add_body(
    'Age: 22-28, recently graduated from university, lives in a regional city. '
    'Comfortable with technology, uses smartphone daily, prefers digital '
    'channels over physical visits. Needs to register as unemployed and '
    'access job search services. Expects fast processing and real-time status '
    'updates. Pain point: traveling to the nearest regional office takes a full '
    'day and costs significant money.'
)

add_heading_3('Persona 2: Mid-Career Worker')
add_body(
    'Age: 35-45, recently lost their job, has family responsibilities. '
    'Moderate technology skills, uses smartphone but may need guidance. Needs '
    'unemployment benefits and retraining program information. Expects clear '
    'communication about eligibility and next steps. Pain point: does not '
    'understand which documents are required and often submits incomplete '
    'applications.'
)

add_heading_3('Persona 3: Employer')
add_body(
    'Small or medium business owner, needs to post job vacancies and find '
    'suitable candidates. Prefers online platforms, expects quick response '
    'from the employment agency. Pain point: current process requires '
    'in-person visits to submit vacancy information, which is time-consuming '
    'for a busy business owner.'
)

doc.add_paragraph()

# ======================================================================
# 5. AS-IS SERVICE JOURNEY ANALYSIS
# ======================================================================
add_heading_1('5. As-Is Service Journey Analysis')

add_heading_2('5.1 Current Process Overview')
add_body(
    'The current employment service process is primarily paper-based and '
    'requires citizens to visit a regional office of the State Employment '
    'Agency. The process involves multiple manual steps, paper document '
    'handling and limited communication between the citizen and the agency '
    'during processing. There is no digital channel for application submission '
    'or status tracking.'
)

add_heading_2('5.2 Current Service Journey Stages')

journey_as_is = [
    ('1. Awareness', 'Citizen learns about available services through word-of-mouth or media', 'Low awareness, no centralized information'),
    ('2. Document Preparation', 'Citizen gathers required documents (ID, diploma, references)', 'Unclear requirements, frequent rejections due to incomplete documents'),
    ('3. Office Visit', 'Citizen travels to regional office, waits in queue, submits paper application', 'Long travel time, queues, limited office hours'),
    ('4. Application Receipt', 'Officer reviews documents, creates paper file, provides receipt number', 'No digital record, manual data entry errors'),
    ('5. Processing', 'Application is reviewed by multiple officers across different departments', 'Average 14 days, no visibility for citizen'),
    ('6. Decision', 'Final decision is made and recorded on paper file', 'No notification system, citizen must call or visit to check'),
    ('7. Result Delivery', 'Citizen visits office again to receive the decision document', 'Second visit required, additional cost and time'),
]
add_simple_table(['Stage', 'Activity', 'Pain Point'], journey_as_is)


doc.add_paragraph()

add_heading_2('5.3 Gap Analysis Summary')

gaps = [
    ('No digital channel', 'All services require physical office visit', 'Online application submission via Telegram bot and web portal'),
    ('No status tracking', 'Citizens cannot check application progress', 'Real-time status updates through digital channels'),
    ('Manual processing', 'Paper-based workflow with manual handoffs', 'Automated workflow with digital document management'),
    ('No integration', 'Data not shared between government organizations', 'Multi-agency data exchange through middleware layer'),
    ('No SLA monitoring', 'Processing time not measured or tracked', 'Real-time SLA and KPI monitoring dashboard'),
]
add_simple_table(['Current Gap', 'Description', 'Proposed Solution'], gaps)

doc.add_paragraph()

# ======================================================================
# 6. TO-BE SERVICE DESIGN
# ======================================================================
add_heading_1('6. To-Be Service Design')

add_heading_2('6.1 Designed Service Journey')
add_body(
    'The redesigned service journey follows a life-event-based approach where '
    'the entire employment service is structured around the citizen need of '
    'finding employment or managing labor-related requirements. The journey '
    'covers the full end-to-end process from initial awareness through final '
    'result delivery, with digital channels as the primary interaction point.'
)

journey_to_be = [
    ('1. Awareness', 'Citizen discovers services through my.gov.az portal, ASAN Service or social media', 'Centralized digital information hub'),
    ('2. Digital Application', 'Citizen submits application via Telegram bot or web portal with guided document checklist', 'Digital form with real-time validation'),
    ('3. Automatic Verification', 'System automatically verifies citizen data with integrated government databases', 'Reduced document requirements, faster processing'),
    ('4. Digital Processing', 'Application is processed through automated workflow with digital routing', 'Target: 5 business days'),
    ('5. Real-Time Tracking', 'Citizen receives status notifications at each processing stage', 'Telegram notifications + portal dashboard'),
    ('6. Digital Decision', 'Decision is recorded digitally and citizen is notified immediately', 'Push notification with result and next steps'),
    ('7. Digital Delivery', 'Citizen receives the decision document through the digital channel', 'PDF download or delivery via ASAN Service'),
]
add_simple_table(['Stage', 'Activity', 'Improvement'], journey_to_be)


doc.add_paragraph()

add_heading_2('6.2 Service Principles')

principles = [
    ('Citizen-first design', 'Every design decision starts from the citizen need, not institutional convenience'),
    ('Digital by default', 'Digital channels are the primary service delivery method, physical visits are the exception'),
    ('Life-event grouping', 'Services are organized around citizen life events, not government departments'),
    ('Transparency', 'Citizens have full visibility into their application status at all times'),
    ('Multi-agency coordination', 'Service delivery requires seamless data exchange between government organizations'),
    ('Continuous improvement', 'Service performance is measured through KPIs and user feedback drives improvements'),
]
add_simple_table(['Principle', 'Description'], principles)

doc.add_paragraph()

# ======================================================================
# 7. SERVICE BLUEPRINT
# ======================================================================
add_heading_1('7. Service Blueprint')

add_heading_2('7.1 Blueprint Overview')
add_body(
    'The service blueprint maps the relationship between the citizen journey '
    'and the back-stage processes, support systems and organizational '
    'structures required to deliver the service. The blueprint identifies '
    'critical touchpoints, handoff points between departments and integration '
    'requirements with external systems.'
)

add_heading_2('7.2 Blueprint Layers')

blueprint = [
    ('Physical Evidence', 'Telegram bot interface, Web portal, Notification messages, PDF decision documents'),
    ('Citizen Actions', 'Open app, Fill application, Upload documents, Check status, Download result'),
    ('Front-Stage (Visible)', 'Welcome message, Application form, Document validation, Status page, Result delivery'),
    ('Back-Stage (Internal)', 'Data verification, Routing to department, Review process, Decision recording, Notification trigger'),
    ('Support Systems', 'Database, Integration middleware, Workflow engine, Monitoring dashboard, Document management'),
]
add_simple_table(['Layer', 'Components'], blueprint)

doc.add_paragraph()

add_heading_2('7.3 Critical Touchpoints')

touchpoints = [
    ('TP-01', 'Application form', 'Digital form with guided document checklist and real-time validation', 'Telegram bot + Web portal'),
    ('TP-02', 'Document upload', 'Digital document upload with format and size validation', 'Web portal'),
    ('TP-03', 'Status notification', 'Push notification at each stage change', 'Telegram bot'),
    ('TP-04', 'Support chat', 'In-app chat for citizen questions and guidance', 'Telegram bot'),
    ('TP-05', 'Result delivery', 'Digital decision document with QR verification code', 'Portal + ASAN delivery'),
]
add_simple_table(['ID', 'Touchpoint', 'Description', 'Channel'], touchpoints)

doc.add_paragraph()

# ======================================================================
# 8. SERVICE CHANNELS AND TOUCHPOINTS
# ======================================================================
add_heading_1('8. Service Channels and Touchpoints')

add_heading_2('8.1 Channel Strategy')
add_body(
    'The service is delivered through a multi-channel approach with the '
    'Telegram bot as the primary channel for application submission and '
    'status tracking. The web portal serves as a secondary channel for more '
    'complex interactions such as document upload and detailed service '
    'information. Physical office visits are reserved for exceptional cases '
    'that cannot be handled digitally.'
)

channels = [
    ('Telegram Bot', 'Primary', 'Application submission, status tracking, notifications, citizen support', 'Target: 70% of interactions'),
    ('Web Portal', 'Secondary', 'Detailed service information, document upload, account management', 'Target: 30% of interactions'),
    ('ASAN Service Centers', 'Exception', 'Cases requiring physical document verification or citizen without digital access', 'Target: Less than 10%'),
    ('Regional Offices', 'Exception', 'Complex cases requiring face-to-face consultation', 'Target: Less than 5%'),
]
add_simple_table(['Channel', 'Role', 'Services', 'Target Volume'], channels)

doc.add_paragraph()

add_heading_2('8.2 Channel Integration')
add_body(
    'All channels share a unified backend system, ensuring that citizen data '
    'and application status are consistent across channels. A citizen can start '
    'an application on the web portal and continue tracking it via the Telegram '
    'bot without any data loss. The integration layer uses a single-window '
    'architecture that connects all channels to the central LMAS system.'
)

# ======================================================================
# 9. MULTI-AGENCY COORDINATION MODEL
# ======================================================================
add_heading_1('9. Multi-Agency Coordination Model')

add_heading_2('9.1 Coordination Framework')
add_body(
    'The LMAS service requires data exchange and coordination with multiple '
    'government organizations. The coordination model defines how agencies '
    'interact, what data is shared and what the responsibilities of each '
    'agency are in the service delivery process. This model ensures that '
    'citizens receive a seamless service experience without needing to '
    'understand the internal organizational structure of government.'
)

add_heading_2('9.2 Agency Roles and Data Exchange')

agency_roles = [
    ('State Employment Agency', 'Service owner', 'Application data, processing status, decision results', 'Provides and consumes'),
    ('Ministry of Labour', 'Policy oversight', 'Employment policies, regulatory requirements', 'Provides'),
    ('Social Protection Fund', 'Data provider', 'Social insurance records, benefit payment history', 'Provides'),
    ('Ministry of Education', 'Data provider', 'Education records, diploma verification', 'Provides'),
    ('State Statistics Committee', 'Data consumer', 'Aggregated employment statistics', 'Consumes'),
    ('Innovation Agency', 'Standards body', 'Digital service standards, methodology guidelines', 'Provides'),
]
add_simple_table(['Agency', 'Role', 'Data', 'Direction'], agency_roles)

doc.add_paragraph()

add_heading_2('9.3 Integration Architecture')
add_body(
    'The integration architecture uses a middleware layer that connects the '
    'LMAS system with external government databases through standardized API '
    'interfaces. The middleware handles data transformation, authentication, '
    'logging and error handling. This approach allows each agency to maintain '
    'its own systems while sharing data through secure, controlled interfaces. '
    'The middleware design is based on the single-window government portal '
    'principle, where the citizen interacts with one service while the system '
    'coordinates with multiple agencies in the background.'
)

# ======================================================================
# 10. SERVICE STANDARDS (SLA AND KPI)
# ======================================================================
add_heading_1('10. Service Standards (SLA and KPI)')

add_heading_2('10.1 Service Level Agreements')

slas = [
    ('SLA-01', 'Application acknowledgment', 'Immediate (automated)', '100%', 'System sends confirmation within 30 seconds'),
    ('SLA-02', 'Document verification', '1 business day', '95%', 'Automatic verification with government databases'),
    ('SLA-03', 'Application processing', '5 business days', '90%', 'Full review and decision by assigned officer'),
    ('SLA-04', 'Status notification', 'Within 2 hours of status change', '99%', 'Automated push notification via Telegram'),
    ('SLA-05', 'Decision delivery', '1 business day after decision', '95%', 'Digital delivery with QR verification'),
    ('SLA-06', 'System availability', '99.5% uptime', '99.5%', 'Planned maintenance outside business hours'),
]
add_simple_table(['ID', 'Service', 'Target', 'Compliance', 'Notes'], slas)

doc.add_paragraph()

add_heading_2('10.2 Key Performance Indicators')

kpis = [
    ('KPI-01', 'Digital channel adoption', 'Percentage of citizens using digital channels', '60%', 'Monthly'),
    ('KPI-02', 'First-contact resolution', 'Citizen issue resolved in first interaction', '80%', 'Monthly'),
    ('KPI-03', 'Application completion rate', 'Applications submitted without abandonment', '85%', 'Weekly'),
    ('KPI-04', 'Average processing time', 'Mean days from submission to decision', '5 days', 'Weekly'),
    ('KPI-05', 'Citizen satisfaction (CSAT)', 'Post-service survey average rating', '4.0/5.0', 'Monthly'),
    ('KPI-06', 'SLA compliance', 'Percentage of SLAs met within target', '90%', 'Monthly'),
    ('KPI-07', 'Error rate', 'Percentage of applications with processing errors', 'Less than 3%', 'Weekly'),
]
add_simple_table(['ID', 'KPI', 'Description', 'Target', 'Frequency'], kpis)

doc.add_paragraph()

add_heading_2('10.3 Monitoring Dashboard')
add_body(
    'A real-time monitoring dashboard is designed for the management board to '
    'track all SLA and KPI indicators. The dashboard provides visual displays '
    'of application volumes, processing times, SLA compliance rates and citizen '
    'satisfaction scores. The dashboard is updated in real-time and accessible '
    'through a web interface. Alerts are configured for SLA breaches and '
    'unusual patterns in application volumes or processing times.'
)

# ======================================================================
# 11. SERVICE PORTFOLIO ALIGNMENT
# ======================================================================
add_heading_1('11. Service Portfolio Alignment')

add_heading_2('11.1 Portfolio Context')
add_body(
    'The LMAS is part of the broader government service portfolio managed by '
    'the State Employment Agency. The service portfolio includes employment '
    'registration, job search, vocational training, labor contract management '
    'and unemployment benefits. The LMAS design ensures alignment with the '
    'existing portfolio structure and uses consistent design patterns, data '
    'models and integration standards across all services.'
)

add_heading_2('11.2 Alignment with Government Service Standards')
add_body(
    'The service design follows the digital government service standards set '
    'by the Innovation and Digital Development Agency. These standards cover '
    'service design methodology, data exchange protocols, security requirements '
    'and citizen experience guidelines. The LMAS is designed to be compatible '
    'with the my.gov.az unified government portal, enabling citizens to access '
    'employment services alongside other government services through a single '
    'entry point.'
)

portfolio = [
    ('Service design methodology', 'Innovation Agency standards', 'Aligned - uses As-Is / To-Be analysis, citizen journey mapping'),
    ('Data exchange', 'Government interoperability framework', 'Aligned - uses standardized API interfaces'),
    ('Authentication', 'my.gov.az single sign-on', 'Aligned - supports SSO integration'),
    ('Notification', 'Government notification standards', 'Aligned - uses official notification channels'),
    ('Accessibility', 'WCA 2.1 AA standards', 'Partially aligned - planned for Phase 2'),
]
add_simple_table(['Area', 'Standard', 'Status'], portfolio)

doc.add_paragraph()

# ======================================================================
# 12. PROCESS ARCHITECTURE
# ======================================================================
add_heading_1('12. Process Architecture')

add_heading_2('12.1 Architecture Overview')
add_body(
    'The process architecture defines the structure, layers and relationships '
    'of all processes within the LMAS. The architecture is designed at three '
    'levels: strategic (service portfolio and life-event grouping), tactical '
    '(end-to-end service processes) and operational (individual task workflows). '
    'This multi-level approach ensures that the service is managed both at the '
    'process architecture level and at the operational execution level, as '
    'required by the Agency standards.'
)

add_heading_2('12.2 Process Layers')

process_layers = [
    ('Strategic layer', 'Service portfolio management, life-event categorization, service priority definition'),
    ('Tactical layer', 'End-to-end service processes (application to result), cross-agency coordination workflows'),
    ('Operational layer', 'Individual task workflows, document verification steps, notification triggers'),
]
add_simple_table(['Layer', 'Scope'], process_layers)

doc.add_paragraph()

add_heading_2('12.3 As-Is / To-Be Process Mapping')
add_body(
    'Each service process has been mapped from As-Is state to To-Be state '
    'using BPMN process models. The As-Is models capture the current paper-based '
    'processes with all manual steps, bottlenecks and handoff points. The To-Be '
    'models define the redesigned digital processes with automated workflows, '
    'validation rules and integration touchpoints. The gap between As-Is and '
    'To-Be has been analyzed and documented, and the transition plan is '
    'included in the Implementation Roadmap section of this document.'
)

# ======================================================================
# 13. IMPLEMENTATION ROADMAP
# ======================================================================
add_heading_1('13. Implementation Roadmap')

add_heading_2('13.1 Phased Approach')
add_body(
    'The implementation follows a phased approach to manage risk and deliver '
    'value incrementally. Each phase delivers a working subset of the service '
    'that can be tested and validated before moving to the next phase. The '
    'roadmap is aligned with the Agency quarterly planning cycle.'
)

roadmap = [
    ('Phase 1', 'Foundation', 'Month 1-3', 'Core system setup, database, integration middleware, Telegram bot MVP'),
    ('Phase 2', 'Core Services', 'Month 4-6', 'Digital application submission, automatic verification, status tracking'),
    ('Phase 3', 'Full Integration', 'Month 7-9', 'Multi-agency data exchange, monitoring dashboard, web portal'),
    ('Phase 4', 'Optimization', 'Month 10-12', 'Performance optimization, accessibility improvements, advanced analytics'),
]
add_simple_table(['Phase', 'Name', 'Timeline', 'Deliverables'], roadmap)

doc.add_paragraph()

add_heading_2('13.2 Critical Dependencies')
deps = [
    ('Government database access', 'API access to Social Protection Fund and Ministry of Education databases', 'High', 'Phase 2'),
    ('my.gov.az integration', 'Single sign-on and unified portal integration', 'Medium', 'Phase 3'),
    ('ASAN Service integration', 'Physical document delivery channel integration', 'Low', 'Phase 3'),
    ('Staff training', 'Training for agency staff on new digital processes', 'High', 'Phase 2'),
    ('Legal framework', 'Regulatory approval for digital application submission', 'High', 'Phase 1'),
]
add_simple_table(['Dependency', 'Description', 'Impact', 'Required By'], deps)

doc.add_paragraph()

# ======================================================================
# 14. RISKS AND DEPENDENCIES
# ======================================================================
add_heading_1('14. Risks and Dependencies')

add_heading_2('14.1 Risk Register')

risks = [
    ('R-01', 'Low citizen adoption of digital channels', 'High', 'Medium', 'User awareness campaign, simplified onboarding, ASAN support'),
    ('R-02', 'Integration delays with government databases', 'High', 'High', 'Early engagement with IT teams, fallback manual process'),
    ('R-03', 'Data quality issues in existing databases', 'Medium', 'High', 'Data cleansing before migration, validation rules'),
    ('R-04', 'Staff resistance to new processes', 'Medium', 'Medium', 'Change management program, training, phased rollout'),
    ('R-05', 'System performance under high load', 'Medium', 'High', 'Load testing, scalable architecture, caching strategy'),
    ('R-06', 'Regulatory changes during implementation', 'Low', 'Medium', 'Regular policy review, flexible process design'),
]
add_simple_table(['ID', 'Risk', 'Probability', 'Impact', 'Mitigation'], risks)

doc.add_paragraph()

# ======================================================================
# 15. APPENDIX: GLOSSARY
# ======================================================================
add_heading_1('15. Appendix: Glossary')

glossary = [
    ('LMAS', 'Labour and Employment Subsystem - the digital employment service system'),
    ('SLA', 'Service Level Agreement - defined performance targets for service delivery'),
    ('KPI', 'Key Performance Indicator - measurable value that shows service performance'),
    ('As-Is', 'Current state of a process before redesign'),
    ('To-Be', 'Future state of a process after redesign'),
    ('BPMN', 'Business Process Model and Notation - standard for process modeling'),
    ('Life-event', 'A significant event in citizen life around which services are grouped'),
    ('Single-window', 'A unified access point for multiple government services'),
    ('Touchpoint', 'Any interaction point between the citizen and the service'),
    ('Service Blueprint', 'A diagram that shows the relationship between service components'),
    ('Service Portfolio', 'The complete set of services offered by an organization'),
    ('Process Architecture', 'The structural design of processes across an organization'),
]
add_simple_table(['Term', 'Definition'], glossary)

# Save
output = '/home/z/my-project/ba-practice/LMAS_Service_Design_Document_SDD.docx'
doc.save(output)
print(f'SDD saved: {output}')
