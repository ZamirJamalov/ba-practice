#!/usr/bin/env python3
"""Generate BA Service Design Interview Guide - comprehensive learning doc for Innovation Agency interview prep"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# ── Helpers ──
DARK = RGBColor(33, 37, 41)
ACCENT = RGBColor(0, 90, 156)
BLUE2 = RGBColor(0, 60, 120)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = RGBColor(120, 120, 120)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level == 1 else BLUE2
        run.font.name = 'Calibri'
    return h

def add_para(text, bold=False, size=10, color=DARK, space_after=Pt(4), space_before=Pt(0),
             alignment=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = color
    return p

def add_bullet(text, level=0, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        r.font.color.rgb = DARK
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        r.font.color.rgb = DARK
    else:
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
        if not p.runs:
            r = p.add_run(text)
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
            r.font.color.rgb = DARK
        else:
            p.runs[0].text = text
    return p

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = header
        r.font.size = Pt(9)
        r.font.name = 'Calibri'
        r.font.color.rgb = ACCENT if header else DARK
    return row

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): 'CCCCCC'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def example_box(title, content):
    """Simulate a highlighted box using indented paragraph with different formatting"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    r = p.add_run(title)
    r.bold = True
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.name = 'Calibri'
    r.font.color.rgb = ACCENT
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.right_indent = Inches(0.3)
    r2 = p2.add_run(content)
    r2.font.size = Pt(9)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = DARK


# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('BA SERVICE DESIGN')
r.bold = True
r.font.size = Pt(26)
r.font.name = 'Calibri'
r.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MÜSAHİBƏ HAZIRLIQ GİDİ')
r.bold = True
r.font.size = Pt(18)
r.font.name = 'Calibri'
r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
r = p.add_run('İnnovasiya və Rəqəmsal İnkişaf Agentliyi\n"Dövlət xidmətlərinin dizaynı və təkmilləşdirilməsi üzrə Baş Biznes Analitik"')
r.font.size = Pt(11)
r.font.name = 'Calibri'
r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run('6 may 2026')
r.font.size = Pt(10)
r.font.name = 'Calibri'
r.font.color.rgb = LIGHT_GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════
heading('MÜNDƏRİCAT', level=1)
toc_items = [
    'Hissə 1: Əsas Anlayışlar — Service Design nədir?',
    'Hissə 2: Müştəri Səyahəti Xəritəsi (Customer Journey Map)',
    'Hissə 3: Xidmət Çertyojı (Service Blueprint)',
    'Hissə 4: Design Thinking',
    'Hissə 5: Həyat Hadisələrinə Əsaslanan Xidmət Dizaynı',
    'Hissə 6: As-Is / To-Be Xidmət Analizi',
    'Hissə 7: Sənin Təcrübənin Service Design Dili ilə İfadəsi',
    'Hissə 8: Müsahibə Sualları və Cavab Şablonları (20 sual)',
    'Hissə 9: Çətin Sualların Strateji Cavabları',
    'Hissə 10: Alətlər — Figma, Miro, Bizagi',
    'Hissə 11: Tez Xatırlama Çeklisti',
]
for item in toc_items:
    add_para(item, size=10, color=DARK, space_after=Pt(3))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART 1: CORE CONCEPTS
# ══════════════════════════════════════════════════════════════════════
heading('HİSSƏ 1: ƏSAS ANLAYIŞLAR — SERVICE DESIGN NƏDİR?', level=1)

heading('1.1 Service Design-in Tərifi', level=2)
add_para(
    'Service Design (Xidmət Dizaynı) — insanların ehtiyaclarına yönəlmiş xidmətlərin '
    'planlaşdırılması, təşkili və təkmilləşdirilməsi prosesidir. Bu, yalnız rəqəmsal '
    'platforma və ya fiziki məhsul dizaynı deyil, xidmətin bütün komponentlərinin '
    'vətəndaş (istifadəçi) baxımından optimized olmasını nəzərdə tutur. Service Design '
    'xidmətin hər bir toxunuş nöqtəsini (touchpoint) — veb-saytdan tutmuş ofisə gedənə '
    'qədər — vətəndaşın gözləntilərinə uyğunlaşdırır.'
)
add_para(
    'Dövlət xidmətləri kontekstində Service Design xüsusilə vacibdir, çünki vətəndaşlar '
    'adətən xidməti seçə bilmirlər — onlar dövlət xidmətinə müraciət etmək məcburiyyətindədirlər. '
    'Bu səbəblə, xidmətin keyfiyyəti vətəndaşın həyatını birbaşa təsir edir. Service Design '
    'vasitəsilə dövlət xidmətləri daha əlçatan, şəffaf, effektiv və vətəndaşa yönəlmiş olur.'
)

heading('1.2 Service Design vs Business Analysis', level=2)
add_para(
    'Bu iki sahə oxşar görünə bilər, amma fokusları fərqlidir. Business Analysis daha çox '
    'biznes tələblərinin toplanması, sənədləşməsi və texniki komandaya təhvil verilməsinə '
    'fokuslanır. Service Design isə xidmətin bütün ekosistemini — istifadəçi təcrübəsindən '
    'tutmuş arxa plandakı proseslərə qədər — holistik şəkildə nəzərdən keçirir.'
)

# Comparison table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, txt in enumerate(['Mezar', 'Business Analysis', 'Service Design']):
    hdr[i].text = ''
    r = hdr[i].paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(9); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT

comparisons = [
    ['Fokus', 'Biznes tələbləri, sistem funksionallığı', 'İstifadəçi təcrübəsi, xidmət ekosistemi'],
    ['Çıxış', 'BRD, FRD, SRS, User Stories', 'Journey Map, Service Blueprint, Prototype'],
    ['Sahibkar', 'BA (tələblər sahibi)', 'Service Designer (istifadəçi təcrübəsi sahibi)'],
    ['Metod', 'Stakeholder müsahibələri, tələblər toplama', 'Etnoqrafiya, user research, co-creation'],
    ['Nəticə', 'Sistem düzgün işləyir', 'Xidmət vətəndaşa xoş təsir bağlayır'],
]
for row_data in comparisons:
    add_table_row(table, row_data)

heading('1.3 Xidmət Ekosistemi Nədir?', level=2)
add_para(
    'Xidmət ekosistemi xidmətin çatdırılması üçün lazım olan bütün elementlərin birgə '
    'sistemidir. Bunlara daxildir: vətəndaş (istifadəçi), xidmət verən qurum, texnologiya '
    'platforması, fiziki infrastruktur (ofislər, terminallar), əməkdaşlar, prosedurlar, '
    'qaydalar, və digər xidmət provayderlər. Məsələn, məşğulluq xidməti ekosisteminə '
    'DMA, EMAS sistemi, ASAN xidmət mərkəzləri, e-gov portal, SİA sistemi, iş yerləri '
    'və digər dövlət qurumları daxildir.'
)
add_para(
    'Service Designer bu ekosistemin hər bir elementini başa düşməli, aralarındakı qarşılıqlı '
    'əlaqəni xəritələndirməli və bottlenecks (dar nöqtələr) müəyyən edib aradan qaldırmalıdır. '
    'Bu yanaşma "systems thinking" adlanır — bütününü görüb, hissələri optimallaşdırmaq.'
)

heading('1.4 Front-Stage vs Back-Stage', level=2)
add_para(
    'Service Design-da xidmətin iki əsas zonası var:'
)
add_bullet('Front-Stage (Səhnə önü): ', bold_prefix='')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
r = p.add_run('Front-Stage ')
r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
r = p.add_run(
    '— vətəndaşın birbaşa gördüyü və qarşılıqlı əlaqədə olduğu hissələr. Məsələn: '
    'veb-sayt interfeysi, ASAN mərkəzdəki qəbul countersı, telefon operatoru ilə söhbət, '
    'SMS bildiriş, Telegram bot. Vətəndaş burada xidmətin "üzünü" görür.'
)
r.font.size = Pt(10); r.font.name = 'Calibri'

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
r = p.add_run('Back-Stage (Səhnə arxası): ')
r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
r = p.add_run(
    '— vətəndaşın görmədiyi, amma xidmətin çatdırılması üçün zəruri olan proseslər. '
    'Məsələn: müraciətin emal olunması, bazada yoxlanılması, digər qurumlarla məlumat '
    'mübadiləsi, qərarvermə prosesi, sənədlərin hazırlanması. Back-Stage-nin səmərəli '
    'işləməməsi Front-Stage-da gecikmə və səhv kimi görünür.'
)
r.font.size = Pt(10); r.font.name = 'Calibri'

heading('1.5 Touchpoints (Toxunuş Nöqtələri)', level=2)
add_para(
    'Touchpoint — vətəndaşın xidmətlə qarşılıqlı əlaqədə olduğu hər bir an və ya kanaldır. '
    'Dövlət xidmətləri üçün tipik touchpoints:'
)
tp_items = [
    'Rəqəmsal: e-gov portal, mobil applikasiya, Telegram bot, SMS, email bildiriş',
    'Fiziki: ASAN xidmət mərkəzi, dövlət idarəsinin qəbul ofisi, terminallar',
    'İnsan: operator, məsləhətçi, həkim, müfəttiş',
    'Sənəd: çap edilmiş blank, arayış, sertifikat',
    'Proses: növbə gözləmə, sənəd toplama, imza atma',
]
for item in tp_items:
    add_bullet(item)

add_para(
    'Service Designer bu touchpoints-ı xəritələndirir, hər birində vətəndaşın nə hiss '
    'etdiyini başa düşür və mənfi "moments of truth" (həqiqət anları) müəyyən edib '
    'yaxşlaşdırır. Məsələn, növbədə gözləmə — mənfi moment. SMS ilə nömrə göndərmək — '
    'bu mənfi anı azalda bilər.',
    space_before=Pt(4)
)

divider()

# ══════════════════════════════════════════════════════════════════════
# PART 2: CUSTOMER JOURNEY MAP
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('HİSSƏ 2: MÜŞTƏRİ SƏYAHƏTİ XƏRİTƏSİ (CUSTOMER JOURNEY MAP)', level=1)

heading('2.1 CJM Nədir?', level=2)
add_para(
    'Customer Journey Map (Müştəri Səyahəti Xəritəsi) — istifadəçinin xidmətlə ilk '
    'tanışlığından sonrakı davranışına qədər olan bütün yolun vizual xəritəsidir. Bu '
    'xəritə hər bir mərhələdə istifadəçinin nə etdiyini, nə düşündüyünü, nə hiss etdiyini '
    'və hansı problemlərlə qarşılaşdığını göstərir. CJM xidmətin hansı hissələrində '
    'istifadəçinin xoş təcrübə aldığını, hansılarında isə problem yaşadığını aşkar edir.'
)

heading('2.2 CJM-in Komponentləri', level=2)
components = [
    ('Persona: ', 'Xəritənin əsasə aldığı istifadəçi profili. Kimdir? Nə ehtiyacı var? Hansı səviyyədə texniki bacarığa malikdir? Dövlət xidmətləri üçün "vətəndaş" — yaşına, təhsilinə, rəqəmsal savadlılığına görə fərqli personalar yaradılır.'),
    ('Mərhələlər (Stages): ', 'Səyahətin böyük zaman axını. Məsələn: "Ehtiyacı müəyyən etmə" → "Məlumat axtarmaq" → "Müraciət etmək" → "Gözləmək" → "Nəticə əldə etmək" → "Sonrakı izləmə".'),
    ('Fəaliyyətlər (Actions): ', 'Hər mərhələdə istifadəçinin etdiyi konkret hərəkətlər. Məsələn: "e-gov-a daxil olmaq", "formu doldurmaq", "sənədləri yükləmək".'),
    ('Toxunuş nöqtələri (Touchpoints): ', 'Hər fəaliyyətin bağlı olduğu kanal. Məsələn: "veb-sayt", "mobil app", "ASAN mərkəzi".'),
    ('Emosiyalar (Emotions): ', 'Hər mərhələdə istifadəçinin hissi. Adətən + (müsbət) və - (mənfi) ilə göstərilir. Məsələn: formanı doldurmaq çətindirsə → mənfi; SMS bildiriş alanda → müsbət.'),
    ('Ağrı nöqtələri (Pain Points): ', 'İstifadəçinin qarşılaşdığı problemlər, çətinliklər, maneələr. Məsələn: "3 saat növbədə gözləmək", "sənəd çatışmaması".'),
    ('İmkanlar (Opportunities): ', 'Pain point-ləri aradan qaldırmaq üçün təkliflər. Məsələn: "online sifariş sistemi", "SMS xatırlatma".'),
]
for prefix, text in components:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(prefix)
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.name = 'Calibri'

heading('2.3 CJM-in Yaradılma Addımları', level=2)
steps = [
    'İstifadəçi tədqiqatı aparın: müsahibələr, sorğular, observation (müşahidə)',
    'Personanı (persona) müəyyən edin: hədəf vətəndaşın profili',
    'Mərhələləri müəyyən edin: xidmətin böyük addımları',
    'Hər mərhələdə fəaliyyətləri, touchpoints-ları və emosiyaları qeyd edin',
    'Pain points və opportunities-ları müəyyən edin',
    'Xəritəni vizuallaşdırın: Miro, Figma, və ya Excel ilə',
    'Komanda ilə paylaşın və prioritetləşdirin',
    'Həll təklifləri hazırlayın və həyata keçirin',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'{step}')

heading('2.4 Nümunə: Vətəndaşın DMA-ya Məşğulluq üçün Müraciəti', level=2)
add_para(
    'Aşağıda real nümunə veririk — vətəndaşın Dövlət Məşğulluq Agentliyinə məşğulluq '
    'üzrə müraciət etməsi journey-si. Bu nümunəni müsahibədə istifadə edə bilərsən.',
    italic=True, color=GRAY, size=9.5
)

# Journey table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, txt in enumerate(['Mərhələ', 'Fəaliyyət', 'Touchpoint', 'Emosiya', 'Pain Point']):
    hdr[i].text = ''
    r = hdr[i].paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(8); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT

journey_data = [
    ['Ehtiyacı müəyyən etmə', 'İş axtarır, maaş almaq üçün qeydiyyat lazımdır', 'Dost/ailə məsləhəti', '+/-', 'Məlumat çatışmazlığı — hara müraciət etmək lazımdır?'],
    ['Məlumat toplamaq', 'DMA saytını ziyarət edir, şərtləri oxuyur', 'e-gov / DMA saytı', '+', 'Saytda məlumat qarışıq, asan tapılmır'],
    ['Müraciət etmək', 'Online formu doldurur, sənədləri yükləyir', 'e-gov portal / ASAN', '-/+', 'Formanın bəzi sahələri aydın deyil, sənəd formatı tələbi qeyri-səlis'],
    ['Gözləmə', 'Müraciətin statusunu izləyir', 'SMS / Portal / Telegram', '--', 'Nə vaxt nəticə alınacağı bəlli deyil, neçə gün gözləmək lazımdır'],
    ['Nəticə əldə etmək', 'Qeydiyyat təsdiqlənir, və ya əlavə sənət tələb olunur', 'SMS / Portal', '+/-', 'Əlavə sənət tələbi → yenidən müraciət → dopinq'],
    ['Sonrakı izləmə', 'Vakansiyalara müraciət edir, müsahibələr', 'DMA portal / Telegram', '+', 'Vakansiya siyahısı məhdud, filtrlər çox sadə'],
]
for row_data in journey_data:
    add_table_row(table, row_data)

add_para(
    'Bu xəritənin analizi göstərir ki, ən böyük pain points: (1) məlumat çatışmazlığı, '
    '(2) sənəd tələblərinin qeyri-səclliyi, (3) gözləmə müddətinin bəlli olmaması. '
    'İmkanlar: (1) Telegram bot ilə real vaxt status izləmə, (2) dashboard ilə SLA '
    'monitorinqi, (3) SİA ilə avtomatik sənəd yoxlama.',
    space_before=Pt(6), size=9.5
)

divider()

# ══════════════════════════════════════════════════════════════════════
# PART 3: SERVICE BLUEPRINT
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('HİSSƏ 3: XİDMƏT ÇERTYOJU (SERVICE BLUEPRINT)', level=1)

heading('3.1 Service Blueprint Nədir?', level=2)
add_para(
    'Service Blueprint — xidmətin bütün komponentlərinin (səhnə önü, səhnə arxası, '
    'dəstək prosesləri) birlikdə göstərən vizual diaqramdır. Customer Journey Map-dan '
    'fərqli olaraq, Blueprint həm istifadəçi təcrübəsini, həm də arxa plandakı prosesləri '
    'göstərir. O, həmçinin "line of interaction" (əlaqə xətti) və "line of visibility" '
    '(görünürlük xətti) kimi konsepsiyaları təqdim edir.'
)

heading('3.2 Blueprint-in 5 Qatı (Layers)', level=2)
layers = [
    ('1. Fiziki Sübutlar (Physical Evidence): ', 'İstifadəçinin gördüyü fiziki və ya rəqəmsal elementlər. Məsələn: veb-sayt interfeysi, ASAN mərkəzinin daxili dizaynı, SMS template-i, blank forması.'),
    ('2. İstifadəçi Fəaliyyətləri (Customer Actions): ', 'İstifadəçinin etdiyi addımlar. Məsələn: formu doldurmaq, nömrə almaq, sənəd təqdim etmək.'),
    ('3. Səhnə Önü (Front-Stage Actions): ', 'İstifadəçi ilə birbaşa əlaqədə olan əməkdaşların fəaliyyətləri. Məsələn: operator müraciəti qəbul edir, konsultant məlumat verir. "Line of Interaction" bunlardan keçir.'),
    ('4. Səhnə Arxası (Back-Stage Actions): ', 'İstifadəçi görmür, amma xidmətin çatdırılması üçün zəruri olan proseslər. Məsələn: bazada yoxlanma, digər qurumla məlumat mübadiləsi, rəy yazılması. "Line of Visibility" bunlardan keçir.'),
    ('5. Dəstək Prosesləri (Support Processes): ', 'Bütün sistemi dəstəkləyən infrastruktur və proseslər. Məsələn: IT sistemləri, HR (əməkdaşların təlimi), qanunvericilik çərçivəsi, data qaydaları.'),
]
for prefix, text in layers:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(prefix)
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.name = 'Calibri'

heading('3.3 CJM vs Service Blueprint — Fərq', level=2)
add_para(
    'CJM istifadəçinin BAXIŞ BUCAĞIDAN (outside-in) xidməti təsvir edir. Blueprint isə '
    'xidməti İÇƏRİDƏN (inside-out) göstərir — həm istifadəçi, həm əməkdaşlar, həm də '
    'sistemlər baxımından. Müsahibədə bu fərqi izah edə bilərsən: "CJM vətəndaşın hiss '
    'etdiyi problemləri müəyyən edir, Blueprint isə həmin problemlərin həlli üçün arxa '
    'plandakı proseslərin necə dəyişdirilməli olduğunu göstərir."'
)

heading('3.4 Nümunə: EMAS Məşğulluq Xidmətinin Blueprint-i', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, txt in enumerate(['Qat', 'Vətəndaş Online Müraciəti', 'Vətəndaş ASAN-da Müraciəti', 'Telegram Bot Müraciəti']):
    hdr[i].text = ''
    r = hdr[i].paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(8); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT

bp_data = [
    ['Fiziki Sübut', 'e-gov portal interfeysi', 'ASAN counter, blank', 'Telegram chat interfeysi'],
    ['İstifadəçi əməli', 'Form doldurur, imzalayır', 'Nömrə alır, gözləyir, sənəd verir', 'Bot-a mesaj yazır, data göndərir'],
    ['Səhnə Önü', '— (avtomatik)', 'Operator müraciəti qəbul edir, yoxlayır', 'Bot avtomatik cavab verir, istiqamətləndirir'],
    ['Səhnə Arxası', 'EMAS-da müraciət yaradılır, SİA ilə yoxlanma', 'Operator EMAS-a daxil edir', 'Bot API vasitəsilə EMAS-a göndərir'],
    ['Dəstək', 'e-gov infrastruktur, DB, SİA API', 'EMAS, DB, HR (operator təlimi)', 'Telegram API, EMAS API, server'],
]
for row_data in bp_data:
    add_table_row(table, row_data)

divider()

# ══════════════════════════════════════════════════════════════════════
# PART 4: DESIGN THINKING
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('HİSSƏ 4: DESIGN THINKING', level=1)

heading('4.1 Design Thinking Nədir?', level=2)
add_para(
    'Design Thinking — insanların real ehtiyaclarını başa düşməyə əsaslanan yaradıcılıq '
    'üçlü problem həlli metodu. Bu yanaşma əvvəlcə Stanford d.Hood universitetində '
    'inkişaf etdirilmişdir və hazırda dünyanın ən böyük şirkətləri və dövlət qurumları '
    'tərəfindən istifadə olunur. Design Thinking 5 mərhələdən ibarətdir və iterativ '
    'xarakter daşıyır — yəni mərhələlər bir-birini izləmək məcburi deyil, istənilən '
    'mərhələyə qayıtmaq olar.'
)

heading('4.2 Beş Mərhələ', level=2)

dt_stages = [
    ('1. Empatiya (Empathize)', 
     'İstifadəçinin dünyasını başa düşmək. Müsahibələr, müşahidə, sorğular, və ya '
     'şəxsi təcrübə vasitəsilə vətəndaşın ehtiyaclarını, narahatlıqlarını, gözləntilərini '
     'öyrənmək. "Etdiyi zaman nə hiss edir?" sualı əsas sualdır. Məsələn: DMA-da '
     'çalışarkən vətəndaşların şikayətlərini dinləmişəm, növbədə gözləyən insanları '
     'müşahidə etmişəm — bu empirik məlumatlar Empatiya mərhələsinə aiddir.'),
    ('2. Problemin Müəyyən Edilməsi (Define)', 
     'Toplanmış məlumatlara əsasən əsas problemi aydın şəkildə ifadə etmək. '
     '"Problem statement" yaradılır. Məsələn: "Məşğulluq üçün müraciət edən vətəndaşlar '
     'müraciətin statusunu real vaxtda bilmirlər, bu isə qeyri-müəyyənlik və etibarsızlıq '
     'yaradır." Bu mərhələdə "How Might We" (HMW) sualları hazırlanır: " necə edək ki, '
     'vətəndaşlar müraciət statusunu real vaxtda izləyə bilsinlər?"'),
    ('3. Fikir Yaratma (Ideate)', 
     'Problem üçün mümkün həll yollarını çoxsaylı və yaradıcı şəkildə düşünmək. '
     'Brainstorming, mind mapping, və ya workshop formatında. Heç bir fikir əvvəlcədən '
     'rədd edilmir. Məsələn: dashboard yaradılması, SMS bildiriş, Telegram bot, email '
     'xəbərdarlıq, mobil app notification, ASAN-da info ekranı — bunların hamısı ideyadır.'),
    ('4. Prototip Yaratma (Prototype)', 
     'Seçilmiş ideyanın sadə və ucuz versiyasını yaratmaq. Bu, tam işləyən sistem deyil, '
     'konsepsiyanın sübutudur. Prototip növü: kağız prototip (sketch), wireframe (Figma), '
     'və ya sadə funksional demo. Məsələn: Telegram bot-un sadə versiyasını yaratdım, '
     'vətəndaşlar mesaj göndərir, bot avtomatik cavab verir — bu prototip idi.'),
    ('5. Test Etma (Test)', 
     'Prototipi real istifadəçilərlə sınaqdan keçirmək. Geri əlaqə toplayır, iterasiya '
     'edir. Nəticə əsasən iki istiqamətdə olur: (a) prototip işləyir — tam versiya '
     'hazırlanır, (b) prototip düzəlişə ehtiyac var — Empatiya mərhələsinə qayıdır. '
     'Məsələn: Telegram bot-u bir qrup vətəndaşla sınaqdan keçirdim, əksəriyyəti '
     'müsbət rəy verdi, bəzi UX düzəlişləri etdim.'),
]
for title, text in dt_stages:
    add_para(title, bold=True, size=10.5, color=ACCENT, space_before=Pt(6))
    add_para(text, size=9.5, space_after=Pt(4))

heading('4.3 Design Thinking Dövlət Xidmətləri üçün Nə Üçün Vacibdir?', level=2)
add_para(
    'Dövlət xidmətləri üçün Design Thinking xüsusilə vacibdir, çünki: (1) dövlət '
    'xidmətləri tez-tez bürokratik yanaşma ilə dizayn olunur — vətəndaş deyil, qaydalar '
    'mərkəzdə olur; (2) Service Design ilk növbədə vətəndaşı mərkəzə qoyur; (3) iterativ '
    'yanaşma böyük investisiya etmədən əvvəl ideyanı sınaqdan keçirməyə imkan verir; '
    '(4) multidissiplinar komandaların (BA, designer, developer, domain expert) birlikdə '
    'işləməsini tələb edir — bu da innovasiyanı stimullaşdırır.'
)

divider()

# ══════════════════════════════════════════════════════════════════════
# PART 5: LIFE-EVENT-BASED SERVICE DESIGN
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('HİSSƏ 5: HƏYAT HADİSƏLƏRİNƏ ƏSASLANAN XİDMƏT DİZAYNI', level=1)

heading('5.1 Konsepsiya Nədir?', level=2)
add_para(
    'Life-Event-Based Service Design (Həyat Hadisələrinə Əsaslanan Xidmət Dizaynı) — '
    'dövlət xidmətlərini ayrı-ayrı prosedurlar deyil, vətəndaşın həyatındakı müəyyən '
    'hadisələr ətrafında qruplaşdırmaq konsepsiyasıdır. Ənənəvi yanaşmada vətəndaş '
    'hər bir xidmət üçün ayrı quruma müraciət edir. Life-event yanaşmada isə vətəndaş '
    'bir "həyat hadisəsi" ilə bağlı bütün xidmətləri bir yerdə əldə edir.'
)
add_para(
    'Bu konsepsiya hazırda Estoniya, Danimarka, Birləşmiş Krallıq və digər inkişaf '
    'etmiş ölkələrdə uğurla tətbiq olunur. Azərbaycanda İnnovasiya Agentliyinin '
    'vakansiyasında bu konsepsiya açıq-qeyd göstərilmişdir — ona görə müsahibədə bu '
    'mövzunu bilmək vacibdir.'
)

heading('5.2 Həyat Hadisələri Nümunələri', level=2)
events = [
    ('Uşağın doğulması: ', 'Təbib qeydiyyatı, doğum haqqı, adın verilməsi, şəxsiyyət vəsiqəsi,Validol connected, sığorta, analıq maaşı — bunların hamısı bir hadisə.'),
    ('Öz biznesini qurmaq: ', 'Nizamnamə təsdiqi, VÖEN, bank hesabı, vergi qeydiyyatı, icazələr, əmlak icarəsi, işçi qəbulu.'),
    ('Təhsil almaq: ', 'Universitetə qəbul, təqaüd, yataqxana, tibbi sığorta, tələbə bileti, məzuniyyət sənədi.'),
    ('Məşğulluq: ', 'İş axtarmaq, müraciət, müsahibə, işə qəbul, sosial sığorta, pensiya.'),
    ('Evlənmək: ', 'Nikah qeydiyyatı, ünvan dəyişikliyi, vergi statusu, əmlak bölgüsü.'),
    ('Pensiyaya çıxmaq: ', 'Təqaüd müraciəti, sığorta yoxlaması, tibbi müayinə, pensiya kartı.'),
]
for prefix, text in events:
    add_bullet(f'{prefix}{text}')

heading('5.3 Life-Event Yanaşmasının Üstünlükləri', level=2)
benefits = [
    'Vətəndaş üçün: Bir mərkəzdən bütün xidmətlər — zaman qənaeti, confusion azalır',
    'Dövlət üçün: Xidmətlərin koordinasiyası, təkrar proseslərin aradan qaldırılması',
    'Effektivlik: Proseslərin avtomatlaşdırılması, data paylaşımı qurumlar arasında',
    'Müştəri məmnuniyyəti: Vətəndaş "hansı quruma müraciət etməliyəm" sualını düşünmür',
    'Şəffaflıq: Bütün proses bir xəritədə görünür, bottleneck-lər asan müəyyən olur',
]
for item in benefits:
    add_bullet(item)

heading('5.4 Müsahibədə Necə İzah Etmək Olar?', level=2)
add_para(
    'Əgər "Life-event based service design barədə nə bilirsiniz?" sualı gələrsə, '
    'belə cavab verə bilərsən (özünə uyğunlaşdır):'
)
example_box(
    'Cavab Şablonu:',
    '"Life-event based service design dövlət xidmətlərini prosedurlar deyil, vətəndaşın '
    'həyat hadisələri ətrafında qruplaşdırmaqdır. Ənənəvi yanaşmada vətəndaş hər xidmət '
    'üçün ayrı quruma gedir. Life-event yanaşmasında isə məsələn, uşağın doğulması '
    'hadisəsi ilə bağlı bütün xidmətlər — təbib, qeydiyyat, şəxsiyyət vəsiqəsi, '
    'sığorta — bir platformada təqdim olunur. Bu yanaşma Estoniya və Danimarkada uğurla '
    'tətbiq olunub. Mən də GPP layihəsində oxşar prinsiplərlə çalışmışam — fərqli qurumların '
    'xidmətlərini vahid portalda birləşdirmək konsepsiyası eynidir."'
)

divider()

# ══════════════════════════════════════════════════════════════════════
# PART 6: AS-IS / TO-BE SERVICE ANALYSIS
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('HİSSƏ 6: AS-IS / TO-BE XİDMƏT ANALİZİ', level=1)

heading('6.1 Service Design Kontekstində As-Is / To-Be', level=2)
add_para(
    'As-Is / To-Be təhlili biznes analitikasının əsas alətidir, amma Service Design '
    'kontekstində fərqli bir perspektivdən baxılır. BA olaraq As-Is/To-Be-i daha çox '
    'proseslər və sistemlər baxımından düşünürük. Service Design-da isə bunu istifadəçi '
    'təcrübəsi baxımından genişləndiririk.'
)
add_para(
    'As-Is Xidmət Təhlili: Mövcud xidmətin bütün ekosistemini — vətəndaş təcrübəsi, '
    'front-stage, back-stage, dəstək prosesləri, data axınları — olduğu kimi xəritələndirmək. '
    'Bu, əslində CJM + Service Blueprint-in birləşməsidir.'
)
add_para(
    'To-Be Xidmət Təhlili: İdeal vəziyyəti dizayn etmək — vətəndaşın ehtiyaclarına '
    'uyğun, effektiv, şəffaf, əlçatan xidmət modeli yaratmaq. To-Be-də həmçinin '
    'teknoogiya, proses, insan və siyasət dəyişiklikləri müəyyən edilir.'
)
add_para(
    'Gap Analysis: As-Is ilə To-Be arasındakı fərq müəyyən etmək və bu fərqləri aradan '
    'qaldırmaq üçün konkret addımlar planlaşdırmaq. Bu BA üçün tanış bir konsepsiyadır — '
    'fərq ondadır ki, Service Design-da "gap" yalnız proses deyil, həm də istifadəçi '
    'təcrübəsi, emosiyalar və touchpoints baxımından müəyyən edilir.'
)

heading('6.2 Dövlət Xidməti As-Is / To-Be Nümunəsi', level=2)
add_para('As-Is vəziyyəti (hal-hazırda):', bold=True, size=10)
asis_items = [
    'Vətəndaş məşğulluq üçün DMA-a gedir, növbə alır, gözləyir',
    'Operatorla danışır, sənədləri təqdim edir',
    'Operator EMAS-a daxil edir, müraciət yaradılır',
    'Müraciət 3-5 gün ərzində baxılır',
    'Nəticə SMS və ya telefonla bildirilir',
    'Əlavə sənət tələbi olarsa, vətəndaş yenə gəlməlidir',
]
for item in asis_items:
    add_bullet(item)

add_para('To-Be vəziyyəti (ideal):', bold=True, size=10, space_before=Pt(4))
tobe_items = [
    'Vətəndaş e-gov və ya Telegram vasitəsilə müraciət edir (0 fiziki ziyarət)',
    'Avtomatik SİA ilə yoxlama — əksər sənədlər avtomatik təsdiqlənir',
    'Müraciət 24 saat ərzində baxılır (SLA: 1 iş günü)',
    'Real vaxt status izləmə — vətəndaş hər an bilir ki, müraciəti haradadır',
    'Əlavə sənət tələbi olarsa, online yükləmək kifayətdir',
    'Nəticə e-imza ilə təsdiqlənir, fiziki gəzməyə ehtiyac yoxdur',
]
for item in tobe_items:
    add_bullet(item)

divider()

# ══════════════════════════════════════════════════════════════════════
# PART 7: YOUR EXPERIENCE IN SERVICE DESIGN LANGUAGE
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('HİSSƏ 7: SƏNİN TƏCRÜBƏNİN SERVICE DESIGN DİLİ İLƏ İFADƏSİ', level=1)

add_para(
    'Bu hissə ən vacib hissədir. Burada sənin mövcud təcrübəni Service Design '
    'terminologiyası ilə necə ifadə etmək lazım olduğunu göstərir. Heç bir yeni bacarıq '
    'uydurmurıq — sadəcə etdiyin işləri Service Design dili ilə təsvir edirik. '
    'Müsahibədə bunları istifadə et.',
    italic=True, color=GRAY, size=9.5
)

heading('7.1 GPP Layihəsi (Merkezi Bank)', level=2)
mappings_gpp = [
    ('Etduyin iş: ', '10+ qurumun GPP-ya inteqrasiyası'),
    ('SD dili: ', 'Multi-agency service ecosystem design — müxtəlif dövlət qurumlarının xidmətlərini vahid platformada birləşdirmək'),
    ('SD konsepsiya: ', 'Life-event based — vətəndaş bir ödəniş üçün ayrı-ayrı qurumlara getmir, vahid portaldan həll edir'),
    ('Journey mapping: ', 'Vətəndaşın ödəniş etmək üçün olan yolunu: banka getmək → növbə → ödəniş Et -> GPP saytı → 1 klik → ödəniş tamam'),
]
for prefix, text in mappings_gpp:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(prefix)
    r.bold = True; r.font.size = Pt(9.5); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.name = 'Calibri'

heading('7.2 EMAS Layihəsi (DMA)', level=2)
mappings_emas = [
    ('Etduyin iş: ', 'Tələblər sənədləşməsi, EMAS sisteminin hazırlanması'),
    ('SD dili: ', 'Service blueprint-in yaradılması — vətəndaşın müraciətindən nəticə əldə etməyədək olan xidmətin bütün qatlarının dizaynı'),
    ('Etduyin iş: ', 'Telegram bot dizayn etdin'),
    ('SD dili: ', 'Yeni service channel-in yaradılması — vətəndaşın xidmətə çıxış yolunun genişləndirilməsi, front-stage touchpoint əlavəsi'),
    ('Etduyin iş: ', 'Monitoring dashboard hazırladın'),
    ('SD dili: ', 'Service performance monitoring — xidmət çatdırılma SLA-ları və KPI-larının real vaxt izlənməsi, back-stage visibility'),
]
for prefix, text in mappings_emas:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(prefix)
    r.bold = True; r.font.size = Pt(9.5); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.name = 'Calibri'

heading('7.3 Embafinans / Fintech Təcrübəsi', level=2)
mappings_emb = [
    ('Etduyin iş: ', 'BRD, FRD, SRS yazdın, User Stories hazırladın'),
    ('SD dili: ', 'Requirements engineering — xidmətin hər bir funksionallığının sistemli şəkildə sənədləşməsi, stakeholder ehtiyaclarından service specifications-a keçid'),
    ('Etduyin iş: ', 'UAT əlaqələndirdin'),
    ('SD dili: ', 'Service validation — hazırlanan xidmətin real istifadəçilərlə sınaqdan keçirilməsi, user acceptance testing through service design lens'),
    ('Etduyin iş: ', 'RICE framework ilə prioritetləşdirdin'),
    ('SD dili: ', 'Feature prioritization based on user value — ən çox vətəndaş dəyəri yaradan xidmət komponentlərinin müəyyən edilməsi'),
]
for prefix, text in mappings_emb:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(prefix)
    r.bold = True; r.font.size = Pt(9.5); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.name = 'Calibri'

heading('7.4 Umico İnteqrasiya Təcrübəsi', level=2)
mappings_umico = [
    ('Etduyin iş: ', '25+ partner inteqrasiyası'),
    ('SD dili: ', 'Service ecosystem expansion — ekosistemin service provider şəbəkəsinin genişləndirilməsi, API-based service integration'),
]
for prefix, text in mappings_umico:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(prefix)
    r.bold = True; r.font.size = Pt(9.5); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.name = 'Calibri'

divider()

# ══════════════════════════════════════════════════════════════════════
# PART 8: INTERVIEW QUESTIONS & ANSWER TEMPLATES
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('HİSSƏ 8: MÜSAHİBƏ SUALLARI VƏ CAVAB ŞABLONLARI (20 SUAL)', level=1)

add_para(
    'Aşağıda Innovation Agency müsahibəsində verilə biləcək 20 sual və sənin təcrübənə '
    'uyğun cavab şablonları verilmişdir. Hər cavabı özünə uyğunlaşdır, yəni üst-üstə '
    'düşmə. Əsas struktur saxla, amma sözləri öz dilində de.',
    italic=True, color=GRAY, size=9.5
)

questions = [
    {
        'q': '1. Service Design nədir və BA ilə necə fərqlənir?',
        'a': '"Service Design xidmətin bütün ekosistemini — istifadəçi təcrübəsindən arxa plandakı proseslərə qədər — holistik şəkildə dizayn etməkdir. BA isə daha çox biznes tələblərinin toplanması və sənədləşməsinə fokuslanır. Amma praktikada bu iki rol çox uzaxlaşır. Mən BA kemi çalışarken CJM, Service Blueprint, və Design Thinking prinsiplərini tətbiq edirəm. Məsələn, EMAS layihəsində vətəndaşın müraciətindən nəticə alana qədər olan service journey-sini analiz etdim və Telegram bot ilə yeni touchpoint yaradaraq bu journey-ni yaxşılaşdırdım."'
    },
    {
        'q': '2. Customer Journey Map necə yaradırsınız?',
        'a': '"Əvvəlcə target personanı müəyyən edirəm — kimdir, nə ehtiyacı var, rəqəmsal savadlılığı hansı səviyyədədir. Sonra journey-nin əsas mərhələlərini çıxarıram: məsələn, məşğulluq xidməti üçün — ehtiyacı müəyyən etmə, məlumat toplamaq, müraciət etmək, gözləmək, nəticə əldə etmək. Hər mərhələdə fəaliyyətləri, touchpoints-ları və emosiyaları qeyd edirəm. Ən vacibi — pain points müəyyən etməkdir. DMA-da çalışarkən gördüm ki, ən böyük pain point müraciət statusunun bəlli olmaması idi — buna görə dashboard yaratdım. Bu, praktiki CJM tətbiqi idi."'
    },
    {
        'q': '3. Service Blueprint ilə CJM-in fərqi nədir?',
        'a': '"CJM istifadəçinin baxış bucağından xidməti təsvir edir — outside-in. Blueprint isə xidməti içəridən göstərir — both inside-out and outside-in. CJM-də əsas sual: vətəndaş nə hiss edir? Blueprint-də əsas sual: bütün sistem necə işləyir? Blueprint 5 qatdan ibarətdir: fiziki sübutlar, istifadəçi fəaliyyətləri, səhnə önü, səhnə arxası və dəstək prosesləri. Mən EMAS-da hər üç kanal (online, ASAN, Telegram) üçün blueprint xəritələndirmişəm."'
    },
    {
        'q': '4. Dövlət xidmətlərində Service Design üçün ən çox rast gəlinən problemlər nələrdir?',
        'a': '"Üç əsas problem: (1) Silo mentality — hər qurum öz xidmətini ayrı dizayn edir, vətəndaş bir həyat hadisəsi üçün 5-6 quruma getmək məcburiyyətində qalır. (2) Kanunvericilik məhdudiyyətləri — bəzi proseslər qanunla müəyyən edilib, dəyişdirmək çətindir. (3) Rəqəmsal bölgü — bəzi vətəndaşlar online xidmət istifadə edə bilir, bəziləri yox — digital divide. Mən GPP layihəsində birinci problemi həll etməyə çalışmışam — fərqli qurumları vahid portalda birləşdirməklə. Telegram bot-da isə üçüncü problemi həll etdim — mobil telefonu olan hər kəs istifadə edə bilər."'
    },
    {
        'q': '5. Həyat hadisələrinə əsaslanan xidmət dizaynı barədə nə bilirsiniz?',
        'a': '"Life-event based approach dövlət xidmətlərini prosedurlar deyil, vətəndaşın həyat hadisələri ətrafında qruplaşdırmaqdır. Məsələn, uşağın doğulması hadisəsi ilə bağlı təbib, qeydiyyat, vəsiqə, sığorta — bunların hamısı bir flow-da olmalıdır. Estoniya bu yanaşmanı ən uğurla tətbiq edən ölkədir. Mən də GPP layihəsində oxşar prinsipləri tətbiq etmişəm — fərqli qurumların ödəniş xidmətlərini vahid portalda birləşdirmək, vətəndaşın bir ödəniş üçün bir neçə quruma getməsinin qarşısını almaq."'
    },
    {
        'q': '6. Design Thinking-in 5 mərhələsini izah edin və təcrübənizdən nümunə verin.',
        'a': '"Empathize, Define, Ideate, Prototype, Test. Mən bunu DMA-da Telegram bot layihəsində tətbiq etmişəm. Empathize: vətəndaşların növbədə gözləməkdən narazı olduğunu müşahidə etdim. Define: problem — vətəndaşlar müraciət statusunu bilmirlər və fiziki gəlmək məcburiyyətindədirlər. Ideate: SMS, email, dashboard, Telegram bot — bir neçə variant düşündüm. Prototype: Telegram bot-un sadə versiyasını yaratdım. Test: bir qrup vətəndaşla sınaqdan keçirdim, əksəriyyəti müsbət rəy verdi, bəzi UX düzəlişləri etdim. Bu klassik Design Thinking dövrü idi."'
    },
    {
        'q': '7. Multi-agency coordination təcrübəniz varmı?',
        'a': '"Bəli, iki əsas təcrübəm var. Birincisi, Merkezi Bankda GPP layihəsi — 10-dan çox dövlət orqanının inteqrasiyası. Hər qurumun öz data formatı, proseduru və texniki infrastrukturu var idi. Mən hər qurumla ayrı-ayrılıqda işləyib, data exchange specification hazırladım və middleware vasitəsilə onları vahid sistemə bağladım. İkincisi, DMA-da EMAS layihəsində 15 nəfərlik komandada çalışıb, SİA, ASAN və digər qurumlarla əlaqə saxlamışam. Əsas çətinlik — hər qurumun fərqli prioritetləri və vaxt cədvəlləri olması idi. Həlli — aparıcı təşkilat kimi standartlaşdırılmış communication plan yaratmaq idi."'
    },
    {
        'q': '8. SLA və KPI monitoring təcrübəniz barədə danışın.',
        'a': '"DMA-da idarəetmə şurası üçün real vaxt monitoring dashboardu yaratmışam. Bu dashboard vətəndaş müraciətləri, cavab müddətləri, xidmət çatdırılma KPI-ları və əməkdaş performansını göstərirdi. Məqsəd şəffaflıq idi — rəhbərlik hər an bilirdi ki, hansı xidmət nə qədər vaxt aparır, hansı əməkdaş yüksək performans göstərir. Embafinans-da isə sprint release-lər üçün KPI-lar müəyyən edirəm: defect rate, on-time delivery, UAT pass rate. Hər sprint-dən sonra retrospective keçirirəm və KPI-ları analiz edirəm."'
    },
    {
        'q': '9. End-to-end process architecture nədir və necə yaratırsınız?',
        'a': '"End-to-end process architecture xidmətin vətəndaşın ilk toxunuşundan sonrakı davranışına qədər olan tam prosesin dizaynıdır. Bu, yalnız IT sistemini deyil, insan, proses, texnologiya və siyasət baxımından bütün komponentləri əhatə edir. Yaratmaq üçün: (1) As-Is journey map hazırlayıram, (2) pain points müəyyən edirəm, (3) To-Be blueprint dizayn edirəm, (4) hər qat üçün tələblər müəyyən edirəm, (5) stakeholder-larla razılaşdırıram, (6) iterativ olaraq test edirəm. EMAS-da müraciət qəbulundan nəticə əldə etməyədək olan tam prosesi bu yanaşma ilə dizayn etmişəm."'
    },
    {
        'q': '10. Dövlət qurumlarında stakeholder management necə olur?',
        'a': '"Dövlət qurumlarında stakeholder management korporativ sektordan fərqlidir. Burada iyerarxiya daha sərt, qərarvermə prosesi daha yavaş, amma razılaşma vacibdir. Mənim yanaşmam: (1) əvvəlcə rəsmi səviyyədə qəbul elan edirəm, (2) hər stakeholder-ın ehtiyaclarını və narahatlıqlarını dinləyirəm, (3) onların dili ilə danışırıq — texniki deyil, dəyər və fayda baxımından, (4) razılaşdırılmış communication plan yaradıram, (5) dəyişiklikləri kiçik addımlarla təqdim edirəm. GPP layihəsində bu yanaşma çox faydalı oldu — hər qurum öz maraqlarını gördükdə, əməkdaşlıq asanlaşdı."'
    },
    {
        'q': '11. Nə üçün dövlət sektorundan fintech-ə keçdiniz və indi niyə qayıtmaq istəyirsiniz?',
        'a': '"Dövlət sektorunda işləyərkən sistemlərin inkişafında iştirak edirdim, amma metodologiya baxımından irəliləmək istəyirdim. Fintech sektoruna keçid etdim ki, beynəlxalq standartlara uyğun BA metodologiyasını tam miqyasda öyrənib tətbiq edim — BRD, FRD, SRS, Service Design, CJM, Agile sənədləşmə prosesləri. 4 il ərzində 14-dən çox istehsalat səviyyəsində BA sənədi hazırladım. İndi bu bacarıqlarımla dövlət xidmətlərinin digitallaşdırılmasına töhfə vermək istəyirəm. Dövlət sektoruna dəyər gətirə biləcəyimi düşünürəm."'
    },
    {
        'q': '12. Sizin BA sertifikatınız yoxdur. Bunun haqqında nə deyirsiniz?',
        'a': '"Sertifikat yoxdur, amma 14-dən çox istehsalat səviyyəsində BA sənədim var. Metodologiyanı real layihələrdə öyrənmişəm. Mənim üçün BA nəzəri bilik deyil, hər gün tətbiq etdiyim praktikadır. Mənə görə, bir BA-nın dəyəri onun nəyi bildiyində deyil, nəyi çatdıra bildiyindədir. Mənim istehsalat sənədlərim, interview-də göstərəcəyim metodoloji biliklərim və real layihə təcrübəm bunun sübutudur. Sertifikat planım da var — yaxın gələcəkdə IIBA certifications-ı tamamlamaq niyyətindəyəm."'
    },
    {
        'q': '13. Dövlət xidmətində citizen-facing digital channel necə dizayn edirsiniz?',
        'a': '"Əvvəlcə citizen journey map hazırlayıram — vətəndaşın xidmətlə ilk əlaqəsindən nəticə əldə etməyədək olan yol. Sonra touchpoint-ları müəyyən edirəm — hansi kanallar mövcuddur, hansı çatışmır. DMA-da mən üç kanal işləmişəm: online (e-gov), fiziki (ASAN), və Telegram. Hər kanal üçün service blueprint hazırlayıram — front-stage və back-stage prosesləri. Əsas prinsip — omnichannel approach: vətəndaq istənilən kanaldan müraciət etsə, eyni keyfiyyətdə xidmət alsın. Telegram bot-u dizayn edərkən Empathize mərhələsindən başlamışam — vətəndaşların çətinliklərini başa düşmək üçün onlarla danışmışam."'
    },
    {
        'q': '14. Process digitilazation təcrübəniz barədə danışın.',
        'a': '"Prosesin digitallaşdırılması sadəcə kağız formu online-a keçirmək deyil. Mən 4 addımlı yanaşma istifadə edirəm: (1) As-Is prosesi BPMN ilə modellemək, (2) bottleneck-ləri müəyyən etmək, (3) To-Be prosesi dizayn etmək — sadələşdirmək, avtomatlaşdırmaq, (4) incrementally həyata keçirmək. EMAS-da məşğulluq prosesini kağız-dan digital-ə çevirmişəm. Ən çətin hissə 2-ci addım idi — insanlar adətən mövcud prosesi dəyişmək istəmirlər. Həlli — data ilə sübut etmək: göstərdim ki, avtomatlaşdırma nə qədər vaxt qənaəti edəcək."'
    },
    {
        'q': '15. Conflict situation — stakeholder-lar razılaşa bilmirlər.',
        'a': '"Əvvəlcə hər tərəfin ehtiyaclarını və narahatlıqlarını ayrı-ayrılıqda dinləyirəm. Sonra ortaq dəyəri tapıram — hər kəsin razı ola biləcəyi məqsəd. Məsələn, Embafinans-da iki stakeholder fərqli prioritetlərə sahib idi. Mən RICE framework ilə hər tələbi scorladım və data ilə göstərdim ki, hansı tələb daha çox dəyər yaradır. Data-driven approach çox kömək edir — şəxsi rəy deyil, obyektiv meyarlar. Dövlət qurumlarında isə əlavə olaraq rəsmi razılaşma prosesi və yuxarı səviyyədə təsdiq lazım olur."'
    },
    {
        'q': '16. Agile və Waterfall hansı hallarda istifadə edirsiniz?',
        'a': '"Dövlət xidmətləri layihələrində hibrid yanaşma ən yaxşısıdır. Əsas framework Waterfall ilə başlayır — qanunvericilik, razılaşma, böyük miqyaslı planlaşdırma. Sonra isə development mərhələsində Agile sprintlərə keçir. Məsələn, EMAS-da əvvəlcə ümumi requirements toplandı (Waterfall), sonra hər modul üçün 2-həftəlik sprintlərlə inkişaf etdirildi (Agile). Fintech-da tam Agile istifadə edirəm. Amma dövlət qurumlarında stakeholder-ların hər sprint-də iştirakı çətin ola bilər — ona görə communication plan və demo sessions vacibdir."'
    },
    {
        'q': '17. Service Design-da research metodları hansılardır?',
        'a': '"Əsas metodlar: (1) User interviews — vətəndaşlarla birbaşa danışmaq, (2) Observation — xidmət mərkəzində vətəndaşları müşahidə etmək, (3) Survey — geniş miqyaslı məlumat toplamaq, (4) Data analysis — mövcud data-dan pattern-lər tapmaq, (5) Stakeholder workshops — əməkdaşlarla birgə brainstorming. DMA-da mən birinci və dördüncünü aktiv istifadə etmişəm — vətəndaş şikayətlərini analiz edib, əməkdaşlarla workshoplar keçirmişəm. Həmçinin dashboard-dan gələn data-ları analiz edərək service bottlenecks müəyyən etmişəm."'
    },
    {
        'q': '18. Prototyping təcrübəniz varmı?',
        'a': '"Bəli, DMA-da Telegram bot-un prototipini yaratmışam. Əvvəlcə kağızda sketcch etdim — vətəndaş bot-a nə yazar, bot nə cavab verər. Sonra sadə funksional versiya yaratdım — mesaj yazanda avtomatik cavab, müraciət statusu soruşanda real data göstərir. Bir qrup vətəndaşla sınaqdan keçirdim, feedback topladım və iterasiya etdim. Embafinans-da isə wireframe-lər vasitəsilə dashboard-un prototipini hazırlamışam — stakeholder-lara göstərib, əvvəlki versiyada düzəlişlər etmişəm. Prototyping böyük investisiya etmədən əvvəl ideyanı sübut etməyin ən yaxşı yoludur."'
    },
    {
        'q': '19. Kanunvericilik məhdudiyyətləri ilə necə işləyirsiniz?',
        'a': '"Dövlət xidmətlərində hər şey qanunla bağlıdır — bu, BA üçün constraint kimi işləyir. Mənim yanaşmam: (1) əvvəlcə qanunvericilik çərçivəsini tam başa düşürəm, (2) constraint-ləri müəyyən edirəm — nəyi dəyişmək olar, nəyi yox, (3) constraint-lər daxilində optimal həll tapıram. Məsələn, EMAS-da bəzi proseslər qanunla müəyyən edilirdi və onları dəyişdirmək mümkün deyildi. Amma həmin prosesləri avtomatlaşdırmaq və vətəndaşa daha əlçatan etmək mümkün idi. Telegram bot məhz bu məntiqlə yaradıldı — prosesi dəyişdirmədik, amma çatdırma kanalını genişləndirdik."'
    },
    {
        'q': '20. Bu vəzifədə ilk 90 gününüzə aid planınız nədir?',
        'a': '"İlk 30 gün: Listen & Learn — bütün stakeholder-larla görüş, mövcud xidmətləri başa düş, data və sənədləri analiz et, citizen journey-ləri müəyyən et. 31-60 gün: As-Is Analysis — hazırkı xidmət ekosisteminin xəritəsini çıxart, pain points və bottlenecks müəyyən et, prioritetləş. 61-90 gün: To-Be Proposal — ən vacib 2-3 xidmət üçün To-Be blueprint hazırla, quick wins təklif et, stakeholder-larla razılaşdır. Məqsəd — birinci gündən dəyər yaratmaq, amma əvvəlcə dərin başa düşmək."'
    },
]

for item in questions:
    add_para(item['q'], bold=True, size=10.5, color=ACCENT, space_before=Pt(8), space_after=Pt(2))
    add_para(item['a'], size=9.5, space_after=Pt(4))

divider()

# ══════════════════════════════════════════════════════════════════════
# PART 9: STRATEGIC ANSWERS FOR DIFFICULT QUESTIONS
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('HİSSƏ 9: ÇƏTİN SUALLARIN STRATEJİ CAVABLARI', level=1)

heading('9.1 "Niyə 4 ildir dövlət sektorunda deyilsiniz?"', level=2)
add_para('Strategiya: "Öyrənmək üçün getdim, öyrəndim, qayıtmaq istəyirəm" narrativi.', bold=True, italic=True, color=GRAY, size=9.5)
add_para(
    '"Dövlət sektorunda işləyəndə sistemlərin inkişafında iştirak edirdim, amma metodologiya '
    'baxımından irəliləmək istəyirdim. O vaxt dövlət qurumlarında BA kimi rəsmi rol yox idi — '
    'mən tələbləri toplayırdım, sistem dizayn edirdim, amma bu "business analysis" adlanmırdı. '
    'Fintech-a keçdim ki, beynəlxalq standartları öyrənəm — BRD, FRD, SRS, Service Design, '
    'Agile sənədləşmə. 4 il ərzində 14 istehsalat sənədi hazırladım. İndi bu bacarıqları '
    'dövlət xidmətlərinə gətirmək istəyirəm. Fintech-dan aldığım əsas dərs — metodoloji '
    'disiplin və istifadəçi mərkəzli yanaşma. Bunu dövlət sektoruna tətbiq etmək istəyirəm."'
)

heading('9.2 "Sizin BA kimi 4 il təcrübəniz var, bu çox az deyilmi?"', level=2)
add_para('Strategiya: "Total relevant experience" göstərmək.', bold=True, italic=True, color=GRAY, size=9.5)
add_para(
    '"Rəsmi BA titulum 4 ildir, amma 18 ildir ki, tələbləri toplayıram, sistemləri dizayn '
    'edirəm, inteqrasiya edirəm. Merkezi Bankda GPP inteqrasiyası — bu BA işidir: 10 '
    'qurumun ehtiyaclarını başa düşüb, data spesifikasiyalarını müəyyən edib, sistemləri '
    'birləşdirmək. DMA-da EMAS layihəsi — tələblərin toplanması, sənədləşməsi, '
    'stakeholder əlaqələndirməsi — bunların hamısı BA funksiyalarıdır. Sadəcə o vaxt bu '
    'roldan "Business Analyst" deyil, "developer" və ya "project lead" adlanırdı. '
    'Metodoloji biliklərimi isə son 4 ildə beynəlxalq standartlar səviyyəsində tamamlamışam."'
)

heading('9.3 "Sertifikatınız yoxdur — buna necə baxırsınız?"', level=2)
add_para('Strategiya: Nəticələrə diqqət çəkmək.', bold=True, italic=True, color=GRAY, size=9.5)
add_para(
    '"Sertifikatlar vacibdir və mən də yaxın zamanda IIBA certification almağı planlayıram. '
    'Amma praktiki təcrübəni dəyərləndirməyi xahiş edirəm. 14 istehsalat səviyyəsində sənəd '
    'hazırlamışam. Hər bir sənəd real istifadəçilər üçün, real komandalar tərəfindən '
    'istifadə olunub. Sertifikat nəzəri bilik sübutudur, mənim sənədlərim isə praktiki '
    'bacarığın sübutudur. Mənə görə, CA-nın dəyəri onun nəyi biləcəyində deyil, nəyi '
    'çaldıra bildiyində (deliver) — mənin deliverable-larım bunu sübut edir."'
)

heading('9.4 "Figma və ya Miro istifadə etmisiniz?"', level=2)
add_para('Strategiya: Etiraf et, amma alternativ göstər.', bold=True, italic=True, color=GRAY, size=9.5)
add_para(
    '"Figma və Miro-dan intensiv istifadə etməmişəm, amma Journey Map-ləri və wireframe-ləri '
    'digər alətlərlə yaratmışam. BPMN modellərini professional alətlərlə qururam, wireframe-ləri '
    'kağız və sadə rəqəmsal alətlərlə sketcch edirəm. Figma və Miro-nu 1-2 həftədə mənimsəmək '
    'mümkündür — bunlar intuitiv alətlərdir. Əsas olan Service Design konsepsiyalarını bilməkdir, '
    'alət yalnız vasitədir. Konsepsiyaları bilirəm, aləti tez öyrənəcəm."'
)

heading('9.5 "Dövlət qurumu ilə korporativ fərqi nədir, necə uyğunlaşacaqsınız?"', level=2)
add_para('Strategiya: Hər ikidə təcrübəni göstərmək.', bold=True, italic=True, color=GRAY, size=9.5)
add_para(
    '"Hər iki mühitdə işləmişəm və fərqləri yaxşı bilirəm. Dövlət qurumunda: iyerarxiya '
    'daha sərt, qərarvermə yavaş, lakin razılaşma vacib, formal kommunikasiya, qanunvericilik '
    'constraint-ləri. Korporativdə: daha sürətli, data-driven, Agile, lakin stakeholder '
    'prioritetləri dəyişkən. Mənim güclü tərəfim hər ikisini başa düşməyimdir. Dövlət '
    'sektorunda qaydaları bilirəm — necə rəsmi yazmaq, necə razılaşmaq, necə iyerarxiyada '
    'navigasiya etmək. Fintech-dan isə gətirdiyim sürət, metodoloji disiplin və data-driven '
    'yanaşma dövlət sektoruna dəyər gətirəcək."'
)

divider()

# ══════════════════════════════════════════════════════════════════════
# PART 10: TOOLS
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('HİSSƏ 10: ALƏTLƏR — FIGMA, MIRO, BİZAGİ', level=1)

heading('10.1 Figma', level=2)
add_para(
    'Figma — browser-based interface design alətidir. Service Designer-lar onu wireframe, '
    'prototype və interface dizaynı üçün istifadə edirlər. BA üçün əsas faydası: xidmətin '
    'istifadəçi interfeysini (digital touchpoint) vizual olaraq göstərmək.'
)
add_para('BA üçün vacib Figma bacarıqları:', bold=True, size=10)
figma_skills = [
    'Wireframe yaratmaq — sadə xətlərlə layout sketcch etmək',
    'Low-fidelity prototype — interaktiv, amma sadə versiya',
    'Component library — təkrar istifadə olunan elementlər (düymə, form field)',
    'Collaboration — komanda ilə eyni faylda real vaxt işləmək',
    'Comments və feedback — stakeholder-ların direkt faylda rəy bildirməsi',
]
for item in figma_skills:
    add_bullet(item)
add_para(
    'Müsahibədə istifadə edə bilərsən: "Figma-da intensiv işləmişəm, amma wireframe və '
    'low-fidelity prototype yaratmaq bacarığım var. Journey Map-ləri və Service Blueprint-ləri '
    'Figma-da vizuallaşdırmaq üçün kamilləşdirmək istəyirəm."',
    space_before=Pt(4), italic=True, color=GRAY, size=9.5
)

heading('10.2 Miro', level=2)
add_para(
    'Miro — online collaborative whiteboard alətidir. Service Designer-lar onu Journey Map, '
    'Service Blueprint, mind map, brainstorming və workshop-lar üçün istifadə edirlər. '
    'Əsas üstünlüyü — komanda üzvləri eyni zamanda bir lövhədə işləyə bilərlər.'
)
add_para('BA üçün vacib Miro bacarıqları:', bold=True, size=10)
miro_skills = [
    'Journey Map template-ləri istifadə etmək',
    'Sticky notes, shapes, arrows, connectors ilə vizual xəritə yaratmaq',
    'Workshop facilitation — real vaxt brainstorming',
    'Stakeholder presentation — xəritəni paylaşmaq və izah etmək',
    'Integration — Jira, Confluence ilə əlaqə',
]
for item in miro_skills:
    add_bullet(item)

heading('10.3 Bizagi Modeler', level=2)
add_para(
    'Bizagi Modeler — pulsuz BPMN proses modellemə alətidir. Service Designer-lar və BA-lar '
    'As-Is / To-Be prosesləri BPMN 2.0 formatında modellemək üçün istifadə edirlər. Ən '
    'büyük üstünlüyü — BPMN qaydalarını avtomatik yoxlayır və simulation (simulyasiya) '
    'imkanı verir.'
)
add_para(
    'Bu aləti BA kimi çox yaxşı bilirsən — müsahibədə mütləq bəhs et. Service Design-da '
    'istifadə qaydası: As-Is prosesi Bizagi-də modelleyirəm → bottleneck-ləri müəyyən edirəm '
    '→ To-Be prosesi dizayn edirəm → simulation ilə sınayıram.',
    italic=True, color=GRAY, size=9.5
)

heading('10.4 Digər Faydalı Alətlər', level=2)
tools = [
    ('Lucidchart: ', 'Flowchart, sequence diagram, org chart yaratmaq. Cloud-based, Jira integration.'),
    ('Draw.io (diagrams.net): ', 'Pulsuz, open-source diagram aləti. BPMN, UML, flowchart dəstəkləyir.'),
    ('Canva: ', 'Sadə prezentasiyalar, infographics, social media content üçün.'),
    ('Mural: ', 'Miro-nun alternativi. Workshop və collaborative work üçün.'),
    ('Notion / Confluence: ', 'Sənədləşmə və knowledge management üçün.'),
]
for prefix, text in tools:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(prefix)
    r.bold = True; r.font.size = Pt(9.5); r.font.name = 'Calibri'; r.font.color.rgb = ACCENT
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.name = 'Calibri'

divider()

# ══════════════════════════════════════════════════════════════════════
# PART 11: QUICK MEMORY CHECKLIST
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('HİSSƏ 11: TEZ XATIRLAMA ÇEKLİSTİ', level=1)

add_para(
    'Müsahibədən əvvəl bu 30 bəndi yadda saxla. Hər biri bir sual üçün cavab verə biləcək '
    'açar konsepsiyadır.',
    italic=True, color=GRAY, size=9.5
)

checklist_items = [
    'Service Design = xidmətin bütün ekosistemini istifadəçi baxımından dizayn etmək',
    'CJM = istifadəçinin xidmətlə ilk əlaqəsindən sonrakı davranışına qədər olan yol xəritəsi',
    'Service Blueprint = xidmətin 5 qatı: fiziki sübut, istifadəçi əməli, front-stage, back-stage, dəstək',
    'CJM = outside-in (vətəndaş baxışı), Blueprint = inside-out + outside-in (sistem baxışı)',
    'Design Thinking 5 mərhələ: Empathize → Define → Ideate → Prototype → Test',
    'Life-event based = xidmətləri prosedurlar deyil, həyat hadisələri ətrafında qruplaşdırmaq',
    'Touchpoint = istifadəçinin xidmətlə qarşılıqlı əlaqədə olduğu hər bir an və ya kanal',
    'Front-stage = vətəndaşın gördüyü (interface, operator, SMS)',
    'Back-stage = vətəndaşın görmədiyi (proses, data mübadiləsi, qərarvermə)',
    'Line of Interaction = front-stage ilə istifadəçi arasında sərhəd',
    'Line of Visibility = front-stage ilə back-stage arasında sərhəd',
    'Pain point = istifadəçinin qarşılaşdığı problem, çətinlik və ya maneə',
    'Moment of Truth = xidmət keyfiyyəti barədə istifadəçinin qəti rəy formalaşdırdığı an',
    'Persona = xəritənin hədəf istifadəçi profili (yaş, peşə, ehtiyac, bacarıq)',
    'HMW (How Might We) = problem üçün yaradıcı həll axtarış sualı',
    'Omnichannel = istifadəçi istənilən kanaldan eyni keyfiyyətdə xidmət alsın',
    'Silo mentality = hər qurumun özünü ayrı görməsi, əlaqəsinin olmaması',
    'Service ecosystem = xidməti çatdırmaq üçün lazım olan bütün elementlərin birgə sistemi',
    'As-Is Service Analysis = mövcud xidmətin CJM + Blueprint birləşməsi',
    'To-Be Service Design = ideal vəziyyət — effektiv, əlçatan, vətəndaşa yönəlmiş',
    'Gap Analysis = As-Is ilə To-Be arasındakı fərq — proses + emosiyalar + touchpoints',
    'Co-creation = stakeholder-larla birgə həll yaratmaq — workshop formatı',
    'Prototype = sadə, ucuz versiya — tam sistem deyil, konsepsiyanın sübutu',
    'Iteration = Prototype test → feedback → düzəliş → yenidən test',
    'SLA (Service Level Agreement) = xidmət səviyyəsi razılaşması — vaxt, keyfiyyət standartı',
    'KPI (Key Performance Indicator) = xidmət performansının əsas göstəricisi',
    'Stakeholder mapping = kim təsir edir, kim təsirlənir — power/interest grid',
    'Change management = proses dəyişikliyini insanlara qəbul etdirmək',
    'Digital divide = bəzi vətəndaşlar digital xidmət istifadə edə bilməməsi',
    'Citizen-centric = vətəndaşı mərkəzə qoyan yanaşma — qaydalar yox, insanlar',
]
for i, item in enumerate(checklist_items, 1):
    add_bullet(f'{item}')

add_para(
    'Bu sənəd İnnovasiya və Rəqəmsal İnkişaf Agentliyinin müsahibəsinə hazırlıq '
    'məqsədilə hazırlanmışdır. Əsasən Service Design konsepsiyaları, müsahibə sualları '
    'və cavab şablonları ehtiva edir. Sənin mövcud təcrübən Service Design dili ilə '
    'ifadə edilmişdir.',
    space_before=Pt(12), italic=True, color=GRAY, size=9
)

# ── Save ──
output = '/home/z/my-project/ba-practice/BA_Service_Design_Interview_Guide.docx'
doc.save(output)
print(f'Service Design Guide saved: {output}')
