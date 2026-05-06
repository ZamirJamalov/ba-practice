#!/usr/bin/env python3
"""Generate Terms of Reference (ToR) for LMAS - English B1 level"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

ACCENT = RGBColor(0, 90, 156)
DARK = RGBColor(33, 37, 41)
GRAY = RGBColor(89, 89, 89)
WHITE = RGBColor(255, 255, 255)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = ACCENT
        run.font.name = 'Calibri'
    return h

def add_heading_2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'
    return h

def add_heading_3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'
    return h

def add_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    return p

def add_simple_table(headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.add_row()
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, '005A9C')
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK
    return table


# ======================================================================
# COVER PAGE
# ======================================================================
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run('STATE EMPLOYMENT AGENCY')
run.font.size = Pt(20)
run.font.color.rgb = ACCENT
run.font.name = 'Calibri'
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(20)
run = subtitle.add_run('Terms of Reference (ToR)')
run.font.size = Pt(28)
run.font.color.rgb = DARK
run.font.name = 'Calibri'
run.bold = True

project = doc.add_paragraph()
project.alignment = WD_ALIGN_PARAGRAPH.CENTER
project.paragraph_format.space_after = Pt(30)
run = project.add_run('Labour and Employment Subsystem (LMAS)\nProject Scope, Objectives and Deliverables')
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Document ID', 'ToR-LMAS-2026-001'),
    ('Version', '1.0'),
    ('Date', 'May 2026'),
    ('Author', 'Business Analysis Team'),
    ('Status', 'Approved'),
]
for i, (label, value) in enumerate(meta_data):
    meta_table.rows[i].cells[0].text = label
    meta_table.rows[i].cells[1].text = value
    for j in range(2):
        for p in meta_table.rows[i].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                run.font.color.rgb = DARK
            if j == 0:
                p.runs[0].bold = True

doc.add_page_break()

# ======================================================================
# 1. INTRODUCTION
# ======================================================================
add_heading_1('1. Introduction')

add_heading_2('1.1 Document Purpose')
add_body(
    'This Terms of Reference (ToR) defines the scope, objectives, deliverables, '
    'timeline and governance structure for the Labour and Employment Subsystem '
    '(LMAS) project. The ToR serves as the foundational document that establishes '
    'the project boundaries and provides a shared understanding among all '
    'stakeholders about what the project will deliver, how it will be delivered '
    'and who is responsible for each component.'
)

add_body(
    'This document is intended for the State Employment Agency management, '
    'the Innovation and Digital Development Agency, the project team and all '
    'partner organizations involved in the LMAS project. It should be read in '
    'conjunction with the Business Requirements Document (BRD) and the Service '
    'Design Document (SDD), which provide detailed business and design '
    'specifications for the project.'
)

add_heading_2('1.2 Background')
add_body(
    'The State Employment Agency currently provides employment services '
    'through a network of regional offices across Azerbaijan. Citizens who '
    'need employment-related services such as unemployment registration, job '
    'search assistance or labor contract management must visit these offices '
    'in person. The current process is paper-based, time-consuming and limits '
    'accessibility for citizens outside of major cities.'
)

add_body(
    'In alignment with the national digital government transformation strategy '
    'led by the Innovation and Digital Development Agency, the State Employment '
    'Agency has initiated the LMAS project to digitize its core employment '
    'services. The project aims to create a digital platform that enables '
    'citizens to access employment services through online channels, reduces '
    'processing time and improves overall service quality.'
)

add_heading_2('1.3 Definitions and Abbreviations')

abbrevs = [
    ('ToR', 'Terms of Reference'),
    ('LMAS', 'Labour and Employment Subsystem'),
    ('BA', 'Business Analyst'),
    ('BRD', 'Business Requirements Document'),
    ('FRD', 'Functional Requirements Document'),
    ('SRS', 'Software Requirements Specification'),
    ('SDD', 'Service Design Document'),
    ('SLA', 'Service Level Agreement'),
    ('KPI', 'Key Performance Indicator'),
    ('SDLC', 'Software Development Life Cycle'),
    ('UAT', 'User Acceptance Testing'),
    ('GPP', 'Government Payment Portal'),
    ('SSO', 'Single Sign-On'),
    ('API', 'Application Programming Interface'),
]
add_simple_table(['Abbreviation', 'Definition'], abbrevs)

doc.add_paragraph()

# ======================================================================
# 2. PROJECT OBJECTIVES
# ======================================================================
add_heading_1('2. Project Objectives')

add_heading_2('2.1 Strategic Goal')
add_body(
    'To transform the delivery of employment services in Azerbaijan by creating '
    'a digital platform that provides citizen-centered, life-event-based '
    'employment services through multiple digital channels, integrated with '
    'relevant government databases and aligned with national digital government '
    'standards.'
)

add_heading_2('2.2 Specific Objectives')

objectives = [
    ('OBJ-01', 'Digital service delivery', 'Enable citizens to submit employment service applications through digital channels (Telegram bot, web portal), reducing the need for physical office visits'),
    ('OBJ-02', 'Process digitization', 'Digitize the end-to-end employment service process from application submission to result delivery, replacing the current paper-based workflow'),
    ('OBJ-03', 'Multi-agency integration', 'Establish data exchange connections with 5+ government organizations (Social Protection Fund, Ministry of Education, Ministry of Labour) to enable automatic data verification'),
    ('OBJ-04', 'Real-time monitoring', 'Build a management dashboard that tracks application volumes, processing times, SLA compliance and citizen satisfaction in real time'),
    ('OBJ-05', 'Processing efficiency', 'Reduce average service processing time from 14 calendar days to 5 business days'),
    ('OBJ-06', 'Citizen experience', 'Achieve citizen satisfaction score of 4.0 or higher (out of 5.0) within the first year of operation'),
    ('OBJ-07', 'Service portfolio alignment', 'Ensure the new digital services are aligned with the existing service portfolio and compatible with the my.gov.az unified portal'),
]
add_simple_table(['ID', 'Objective', 'Description'], objectives)

doc.add_paragraph()

add_heading_2('2.3 Success Criteria')

success = [
    ('Digital adoption rate', 'At least 60% of employment service applications submitted through digital channels within 12 months'),
    ('Processing time', 'Average processing time reduced to 5 business days or less'),
    ('SLA compliance', 'At least 90% of applications processed within defined SLA targets'),
    ('System uptime', 'Digital channels available at least 99.5% of the time'),
    ('Citizen satisfaction', 'Average satisfaction score of 4.0 or higher from post-service surveys'),
    ('Multi-agency integration', 'Data exchange operational with at least 3 government organizations at go-live'),
]
add_simple_table(['Criterion', 'Target'], success)

doc.add_paragraph()

# ======================================================================
# 3. PROJECT SCOPE
# ======================================================================
add_heading_1('3. Project Scope')

add_heading_2('3.1 In Scope')

in_scope = [
    ('Digital application submission', 'Citizens can submit employment service applications through Telegram bot and web portal with guided document checklist'),
    ('Automatic data verification', 'System verifies citizen data with integrated government databases (Social Protection Fund, Ministry of Education)'),
    ('Application processing workflow', 'Digital workflow that routes applications to the responsible department, tracks status and records decisions'),
    ('Real-time status tracking', 'Citizens receive notifications at each stage of application processing via Telegram and web portal'),
    ('Monitoring dashboard', 'Management dashboard displaying application volumes, processing times, SLA compliance and KPI indicators'),
    ('Multi-agency data exchange', 'Middleware layer for secure data exchange between LMAS and 5+ government organizations'),
    ('Document management', 'Digital storage and management of all application-related documents'),
    ('Integration with my.gov.az', 'Single sign-on access to employment services through the national government portal'),
]
add_simple_table(['Deliverable', 'Description'], in_scope)

doc.add_paragraph()

add_heading_2('3.2 Out of Scope')

out_scope = [
    ('Mobile native application', 'Development of a dedicated iOS/Android mobile application (future phase)'),
    ('Employer portal', 'Full employer-facing portal for vacancy posting and candidate management (Phase 2)'),
    ('Labor contract management', 'Digital signing and management of labor contracts (Phase 2)'),
    ('Retraining program management', 'Digital management of vocational training program enrollment and tracking (Phase 3)'),
    ('Legacy system migration', 'Migration of historical paper records into the digital system (separate project)'),
    ('Hardware procurement', 'Procurement of servers, network equipment or office hardware (managed separately by IT department)'),
    ('Organizational restructuring', 'Changes to the organizational structure of the State Employment Agency (policy decision)'),
]
add_simple_table(['Item', 'Reason for Exclusion'], out_scope)

doc.add_paragraph()

add_heading_2('3.3 Assumptions')

assumptions = [
    ('A-01', 'API access to government databases will be granted by the Social Protection Fund and Ministry of Education before Phase 2 development starts'),
    ('A-02', 'The Innovation and Digital Development Agency will provide technical standards and methodology guidelines for digital service design'),
    ('A-03', 'The State Employment Agency management will assign a dedicated product owner who has authority to make decisions on business requirements'),
    ('A-04', 'Citizen data required for automatic verification is available in digital format in the relevant government databases'),
    ('A-05', 'The project team will have access to the ASAN Service infrastructure for physical document delivery integration'),
]
add_simple_table(['ID', 'Assumption'], assumptions)

doc.add_paragraph()

add_heading_2('3.4 Constraints')

constraints = [
    ('C-01', 'The system must comply with the data protection and privacy regulations of the Republic of Azerbaijan'),
    ('C-02', 'Development must follow the Agile/Scrum methodology with 2-week sprint cycles'),
    ('C-03', 'All digital channels must support Azerbaijani language (Russian language support is a Phase 2 item)'),
    ('C-04', 'The system must be compatible with the existing government IT infrastructure and security standards'),
    ('C-05', 'Total project budget must not exceed the approved allocation without formal change request'),
    ('C-06', 'Go-live must occur within the 12-month project timeline'),
]
add_simple_table(['ID', 'Constraint'], constraints)

doc.add_paragraph()

# ======================================================================
# 4. DELIVERABLES
# ======================================================================
add_heading_1('4. Deliverables')

add_heading_2('4.1 Business Analysis Deliverables')

ba_deliverables = [
    ('D-01', 'Terms of Reference (ToR)', 'Project scope, objectives, timeline and governance structure', 'Phase 1', 'BA Lead', 'This document'),
    ('D-02', 'Business Requirements Document (BRD)', 'Business objectives, scope, stakeholder analysis and high-level requirements', 'Phase 1', 'BA Lead', 'BRD-LMAS-2026-001'),
    ('D-03', 'Service Design Document (SDD)', 'Service journey, blueprint, channel strategy, SLA/KPI framework', 'Phase 1', 'BA Lead', 'SDD-LMAS-2026-001'),
    ('D-04', 'Functional Requirements Document (FRD)', 'Detailed functional requirements with user stories and acceptance criteria', 'Phase 2', 'BA Lead', 'FRD-LMAS-2026-001'),
    ('D-05', 'Software Requirements Specification (SRS)', 'Technical system requirements, integration specifications and data models', 'Phase 2', 'BA + Tech Lead', 'SRS-LMAS-2026-001'),
    ('D-06', 'API Specifications (Swagger/OpenAPI)', 'REST API endpoint definitions for internal and external integrations', 'Phase 2', 'BA + Dev Team', 'API-LMAS-2026-001'),
    ('D-07', 'As-Is / To-Be Process Models (BPMN)', 'Visual process models for current and future service processes', 'Phase 1', 'BA Lead', 'BPMN-LMAS-2026-001'),
    ('D-08', 'UAT Plan and Test Cases', 'User acceptance testing plan, test scenarios and execution report', 'Phase 3', 'BA Lead + QA', 'UAT-LMAS-2026-001'),
]
add_simple_table(
    ['ID', 'Deliverable', 'Description', 'Phase', 'Responsible', 'Document ID'],
    ba_deliverables
)

doc.add_paragraph()

add_heading_2('4.2 Technical Deliverables')

tech_deliverables = [
    ('D-09', 'System Architecture Document', 'Technical architecture, technology stack, deployment model', 'Phase 1', 'Tech Lead'),
    ('D-10', 'Database Schema Design', 'Data models, entity relationships, migration plan', 'Phase 2', 'Dev Team'),
    ('D-11', 'Integration Middleware', 'Multi-agency data exchange layer with API connectors', 'Phase 2', 'Dev Team'),
    ('D-12', 'Telegram Bot Application', 'Citizen-facing Telegram bot for application submission and tracking', 'Phase 2', 'Dev Team'),
    ('D-13', 'Web Portal', 'Citizen-facing web portal for detailed service access', 'Phase 3', 'Dev Team'),
    ('D-14', 'Monitoring Dashboard', 'Management dashboard for SLA, KPI and application tracking', 'Phase 3', 'Dev Team'),
]
add_simple_table(
    ['ID', 'Deliverable', 'Description', 'Phase', 'Responsible'],
    tech_deliverables
)

doc.add_paragraph()

# ======================================================================
# 5. PROJECT TIMELINE
# ======================================================================
add_heading_1('5. Project Timeline')

add_heading_2('5.1 Phase Overview')

phases = [
    ('Phase 1: Foundation', 'Month 1-3', 'Requirements gathering, ToR/BRD/SDD preparation, system architecture design, development environment setup, vendor selection if needed'),
    ('Phase 2: Core Development', 'Month 4-6', 'Telegram bot development, core application processing workflow, database implementation, integration with Social Protection Fund and Ministry of Education APIs, FRD and SRS finalization'),
    ('Phase 3: Integration and Testing', 'Month 7-9', 'Web portal development, multi-agency integration testing, UAT execution, monitoring dashboard development, my.gov.az SSO integration'),
    ('Phase 4: Go-Live and Optimization', 'Month 10-12', 'User training, pilot launch with limited user group, full go-live, performance optimization, accessibility improvements, post-launch monitoring'),
]
add_simple_table(['Phase', 'Timeline', 'Key Activities'], phases)

doc.add_paragraph()

add_heading_2('5.2 Key Milestones')

milestones = [
    ('M-01', 'ToR approved', 'End of Month 1', 'Project governance established, scope confirmed'),
    ('M-02', 'BRD and SDD approved', 'End of Month 2', 'Business and service design requirements finalized'),
    ('M-03', 'System architecture approved', 'End of Month 3', 'Technical foundation ready for development'),
    ('M-04', 'Telegram bot MVP ready', 'End of Month 5', 'Core digital channel operational for testing'),
    ('M-05', 'Core workflow operational', 'End of Month 6', 'End-to-end application processing functional'),
    ('M-06', 'UAT completed', 'End of Month 9', 'System tested and validated by business stakeholders'),
    ('M-07', 'Pilot launch', 'Month 10', 'Limited user group begins using the system'),
    ('M-08', 'Full go-live', 'Month 12', 'System available to all citizens nationwide'),
]
add_simple_table(['ID', 'Milestone', 'Date', 'Description'], milestones)

doc.add_paragraph()

# ======================================================================
# 6. TEAM STRUCTURE AND ROLES
# ======================================================================
add_heading_1('6. Team Structure and Roles')

add_heading_2('6.1 Project Organization')
add_body(
    'The project team operates under a matrix structure where team members are '
    'assigned from multiple organizations. The State Employment Agency provides '
    'the product owner and business domain experts. The IT department provides '
    'the development team. The Innovation and Digital Development Agency provides '
    'methodology oversight and digital service standards guidance.'
)

add_heading_2('6.2 Roles and Responsibilities')

roles = [
    ('Project Sponsor', 'State Employment Agency management', 'Strategic direction, budget approval, escalation resolution'),
    ('Product Owner', 'State Employment Agency', 'Business requirements prioritization, acceptance decisions, stakeholder alignment'),
    ('BA Lead', 'Business Analysis Team', 'Requirements elicitation, documentation (ToR, BRD, SDD, FRD, SRS), stakeholder communication, UAT coordination'),
    ('Technical Lead', 'IT Department', 'System architecture design, technology selection, technical oversight'),
    ('Development Team', 'IT Department', 'Software development, database implementation, API development, integration'),
    ('QA Specialist', 'IT Department', 'Test planning, test execution, defect tracking'),
    ('UX/UI Designer', 'External / IT Department', 'User interface design, citizen experience optimization'),
    ('Agency Liaison', 'Innovation and Digital Development Agency', 'Methodology guidance, standards compliance review'),
]
add_simple_table(['Role', 'Organization', 'Responsibilities'], roles)

doc.add_paragraph()

add_heading_2('6.3 Governance Structure')

governance = [
    ('Steering Committee', 'Monthly', 'Project Sponsor, Product Owner, BA Lead, Tech Lead, Agency Liaison', 'Strategic decisions, budget review, risk escalation, scope changes'),
    ('Sprint Review', 'Bi-weekly', 'Product Owner, BA Lead, Development Team', 'Sprint demo, backlog review, acceptance decisions'),
    ('Technical Review', 'Weekly', 'Tech Lead, Development Team, BA Lead', 'Technical progress, architecture decisions, integration status'),
    ('Stakeholder Sync', 'Monthly', 'Product Owner, BA Lead, Regional Office representatives', 'User feedback, process improvement, training needs'),
]
add_simple_table(['Forum', 'Frequency', 'Participants', 'Purpose'], governance)

doc.add_paragraph()

# ======================================================================
# 7. STAKEHOLDER COMMUNICATION
# ======================================================================
add_heading_1('7. Stakeholder Communication')

add_heading_2('7.1 Communication Plan')

comm_plan = [
    ('State Employment Agency management', 'Steering Committee', 'Monthly', 'Project status report, risk register, budget status'),
    ('Ministry of Labour', 'Alignment meeting', 'Monthly', 'Policy updates, regulatory changes, progress overview'),
    ('Innovation Agency', 'Progress review', 'Monthly', 'Methodology alignment, standards compliance, design review'),
    ('Product Owner', 'Sprint review', 'Bi-weekly', 'Sprint demo, backlog priorities, acceptance decisions'),
    ('Development Team', 'Stand-up / Sprint planning', 'Daily / Bi-weekly', 'Task progress, blockers, technical decisions'),
    ('Regional Offices', 'Feedback session', 'Monthly', 'User feedback, process issues, training needs'),
    ('Partner agencies', 'Integration sync', 'Monthly', 'API status, data quality, testing coordination'),
]
add_simple_table(['Stakeholder', 'Channel', 'Frequency', 'Content'], comm_plan)

doc.add_paragraph()

add_heading_2('7.2 Reporting')
add_body(
    'The BA Lead is responsible for preparing the following reports: '
    'a monthly project status report covering progress against milestones, '
    'risks and issues; a bi-weekly sprint summary including completed and '
    'planned deliverables; and an ad-hoc issue escalation report when critical '
    'risks or blockers require steering committee attention. All reports follow '
    'the template defined by the Innovation and Digital Development Agency.'
)

# ======================================================================
# 8. RISK MANAGEMENT
# ======================================================================
add_heading_1('8. Risk Management')

add_heading_2('8.1 Risk Register')

risks = [
    ('R-01', 'Delays in government database API access', 'High', 'High', 'Early engagement with partner agencies, fallback manual verification process'),
    ('R-02', 'Scope creep from stakeholder requests', 'High', 'Medium', 'Formal change request process, steering committee approval for scope changes'),
    ('R-03', 'Low digital adoption by citizens', 'Medium', 'High', 'Awareness campaign, simplified onboarding, ASAN Service support channel'),
    ('R-04', 'Data quality issues in source databases', 'Medium', 'High', 'Data cleansing before integration, validation rules, error handling'),
    ('R-05', 'Resource availability constraints', 'Medium', 'Medium', 'Early resource planning, cross-training, contingency allocation'),
    ('R-06', 'Regulatory changes during project', 'Low', 'Medium', 'Regular policy review, flexible design approach, Agency liaison'),
    ('R-07', 'Technology platform instability', 'Low', 'High', 'Architecture review, load testing, scalable infrastructure design'),
]
add_simple_table(['ID', 'Risk', 'Probability', 'Impact', 'Mitigation'], risks)

doc.add_paragraph()

add_heading_2('8.2 Issue Escalation')
add_body(
    'Issues that cannot be resolved at the team level within 5 business days '
    'must be escalated to the steering committee. Critical issues that threaten '
    'project milestones or budget must be escalated within 24 hours. The '
    'escalation path follows: team level, to BA Lead and Tech Lead, to '
    'Product Owner, to Steering Committee, to Project Sponsor.'
)

# ======================================================================
# 9. QUALITY STANDARDS
# ======================================================================
add_heading_1('9. Quality Standards')

add_heading_2('9.1 Documentation Standards')
add_body(
    'All project documents must follow the templates and standards defined by '
    'the Innovation and Digital Development Agency. Documents are written in '
    'English (professional working proficiency level) and use Calibri font for '
    'body text. Each document includes version control, author information, '
    'review history and approval signatures.'
)

add_heading_2('9.2 Development Standards')

dev_standards = [
    ('Code quality', 'Code review required for all pull requests, minimum 80% test coverage'),
    ('API design', 'RESTful API following OpenAPI 3.0 specification'),
    ('Data security', 'Encryption for data in transit (TLS 1.2+) and at rest, role-based access control'),
    ('Accessibility', 'WCA 2.1 AA compliance for web portal (Phase 3 target)'),
    ('Performance', 'API response time under 500ms, page load time under 3 seconds'),
    ('Version control', 'Git-based version control with branching strategy aligned to sprint cycles'),
]
add_simple_table(['Area', 'Standard'], dev_standards)

doc.add_paragraph()

add_heading_2('9.3 Acceptance Criteria')
add_body(
    'Each deliverable must meet the following acceptance criteria before it '
    'is considered complete: the deliverable has been reviewed by the Product '
    'Owner, all feedback has been addressed, the document has been updated to '
    'the latest version, and formal sign-off has been received. For technical '
    'deliverables, UAT must be completed with a pass rate of at least 95% of '
    'test cases before go-live approval.'
)

# ======================================================================
# 10. APPROVAL
# ======================================================================
add_heading_1('10. Approval')

add_body(
    'This Terms of Reference has been reviewed and approved by the following '
    'stakeholders. By signing below, each stakeholder confirms their agreement '
    'with the project scope, objectives, deliverables and timeline as described '
    'in this document.'
)

doc.add_paragraph()

approval = [
    ('Project Sponsor', 'State Employment Agency', '________________', '____/____/________'),
    ('Product Owner', 'State Employment Agency', '________________', '____/____/________'),
    ('BA Lead', 'Business Analysis Team', '________________', '____/____/________'),
    ('Technical Lead', 'IT Department', '________________', '____/____/________'),
    ('Agency Liaison', 'Innovation and Digital Development Agency', '________________', '____/____/________'),
]
add_simple_table(['Role', 'Organization', 'Signature', 'Date'], approval)

doc.add_paragraph()

add_heading_2('10.1 Document Revision History')

revisions = [
    ('1.0', 'May 2026', 'Initial version', 'BA Lead', 'Approved'),
]
add_simple_table(['Version', 'Date', 'Description', 'Author', 'Status'], revisions)

# Save
output = '/home/z/my-project/ba-practice/LMAS_Terms_of_Reference_ToR.docx'
doc.save(output)
print(f'ToR saved: {output}')
